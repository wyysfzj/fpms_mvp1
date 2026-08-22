from __future__ import annotations

import hashlib
import json
import re
import shutil
import subprocess
import zipfile
from dataclasses import asdict, is_dataclass
from pathlib import Path
from typing import Any, Mapping
from uuid import uuid4

import pytest
from docx import Document
from docxtpl import DocxTemplate
from sqlalchemy import select
from sqlalchemy.orm import Session, sessionmaker

from app.modules.templates.models import FormatLetterMapping, Template
from scripts import seed_dev

REPO_ROOT = Path(__file__).resolve().parents[2]
BACKEND_ROOT = REPO_ROOT / "backend"
TEMPLATE_ROOT = BACKEND_ROOT / "storage" / "templates" / "format_letters"
SOURCE_ROOT = REPO_ROOT / "docs" / "postdemo" / "文件样例及模版" / "常用邮件模板"
DATASET_ID = "FPMS-FORMAT-LETTER-CUSTOMER-20260610-V1"
OUTPUT_NAME_RULE = "{case_no}-给{applicant_name}的邮件.docx"
COMMON_VARIABLES = {
    "salutation_text",
    "client_reference_no",
    "case_no",
    "invention_title",
    "application_no",
    "filing_date_text",
    "applicant_names_text",
    "source_notice_name",
}
REPRESENTATIVE_CONTEXT = {
    "salutation_text": "尊敬的客户：您好",
    "client_reference_no": "CLIENT-REF-V8",
    "case_no": "CASE-V8-001",
    "invention_title": "示例发明名称",
    "application_no": "CN202600000001",
    "filing_date_text": "2026-01-02",
    "applicant_names_text": "示例申请人甲；示例申请人乙",
    "source_notice_name": "示例官方通知",
    "publication_date_text": "2026-06-30",
    "deadline_text": "2026-09-30",
    "inventor_names_text": "示例发明人",
    "amount_lines_text": "专利登记费：100.00 CNY\n印花税：5.00 CNY",
    "notice_variant_code": "OA_FIRST",
    "template_variant_code": "FORMAT_LETTER_007",
}

EXPECTED_ROWS = (
    {
        "code": "FORMAT_LETTER_001",
        "customer_format_letter_name": "官文转发-国内客户-驳回通知",
        "official_doc_name_pattern": "驳回决定",
        "source_path": "docs/postdemo/文件样例及模版/常用邮件模板/驳回通知.doc",
        "source_sha256": "4f8f24d83bb3ca84f4663a0c46a1a84fa060ceb5881d2b9d3001fd074e81b4f2",
        "template_version_id": "FPMS-FORMAT-LETTER-CUSTOMER-20260610-V1-001",
        "template_path": "templates/format_letters/format_letter_001.docx",
        "variables": COMMON_VARIABLES | {"deadline_text"},
        "headers": (
            "序号",
            "贵方案号",
            "我方案号",
            "发明创造名称",
            "申请号",
            "申请日",
            "申请人名称",
            "文件名称",
            "复审期限",
        ),
        "cell_variables": {"复审期限": "deadline_text"},
    },
    {
        "code": "FORMAT_LETTER_002",
        "customer_format_letter_name": "官文转发-国内客户-初审合格",
        "official_doc_name_pattern": "初步审查合格",
        "source_path": "docs/postdemo/文件样例及模版/常用邮件模板/初审合格.doc",
        "source_sha256": "979713345eb2d5d8f4ee02421ebf5de6936b9f4418648dcffbe7a71f4cd62724",
        "template_version_id": "FPMS-FORMAT-LETTER-CUSTOMER-20260610-V1-002",
        "template_path": "templates/format_letters/format_letter_002.docx",
        "variables": COMMON_VARIABLES,
        "headers": (
            "序号",
            "贵方案号",
            "我方案号",
            "发明创造名称",
            "申请号",
            "申请日",
            "申请人名称",
            "文件名称",
        ),
        "cell_variables": {},
    },
    {
        "code": "FORMAT_LETTER_003",
        "customer_format_letter_name": "官文转发-国内客户-公布通知",
        "official_doc_name_pattern": "公布通知书",
        "source_path": "docs/postdemo/文件样例及模版/常用邮件模板/公布通知.doc",
        "source_sha256": "01502d4d3329adff358b3dfa0f995c5bcebb831b02155524140fbdf4995a81da",
        "template_version_id": "FPMS-FORMAT-LETTER-CUSTOMER-20260610-V1-003",
        "template_path": "templates/format_letters/format_letter_003.docx",
        "variables": COMMON_VARIABLES | {"publication_date_text", "deadline_text"},
        "headers": (
            "序号",
            "贵方案号",
            "我方案号",
            "发明创造名称",
            "申请号",
            "申请日",
            "申请人名称",
            "文件名称",
            "公开日",
            "实审期限",
        ),
        "cell_variables": {
            "公开日": "publication_date_text",
            "实审期限": "deadline_text",
        },
    },
    {
        "code": "FORMAT_LETTER_004",
        "customer_format_letter_name": "官文转发-国内客户-实审通知",
        "official_doc_name_pattern": "进入实审通知",
        "source_path": "docs/postdemo/文件样例及模版/常用邮件模板/实审通知.doc",
        "source_sha256": "dec38f3f2999b35c39b6d9cfa9f204bd0c930132da750170b1cdf90bd4666c00",
        "template_version_id": "FPMS-FORMAT-LETTER-CUSTOMER-20260610-V1-004",
        "template_path": "templates/format_letters/format_letter_004.docx",
        "variables": COMMON_VARIABLES,
        "headers": (
            "序号",
            "贵方案号",
            "我方案号",
            "发明创造名称",
            "申请号",
            "申请日",
            "申请人名称",
            "文件名称",
        ),
        "cell_variables": {},
    },
    {
        "code": "FORMAT_LETTER_005",
        "customer_format_letter_name": "官文转发-国内客户-受通",
        "official_doc_name_pattern": "受理通知-电子",
        "source_path": "docs/postdemo/文件样例及模版/常用邮件模板/受通.doc",
        "source_sha256": "acc06a13d4b09349d2eb81fadd31509ca7e260047df21ab38d65621d4607fff0",
        "template_version_id": "FPMS-FORMAT-LETTER-CUSTOMER-20260610-V1-005",
        "template_path": "templates/format_letters/format_letter_005.docx",
        "variables": COMMON_VARIABLES | {"inventor_names_text"},
        "headers": (
            "序号",
            "贵方案号",
            "我方案号",
            "发明创造名称",
            "申请号",
            "申请日",
            "发明人",
            "申请人名称",
            "文件名称",
        ),
        "cell_variables": {"发明人": "inventor_names_text"},
    },
    {
        "code": "FORMAT_LETTER_006",
        "customer_format_letter_name": "官文转发-国内客户-授权通知",
        "official_doc_name_pattern": "授权通知书-电子",
        "source_path": "docs/postdemo/文件样例及模版/常用邮件模板/授权通知.doc",
        "source_sha256": "83c14cf4b6514bee2c3f084cb13d2fb66f067dbeb8260de659f9ffb18e542974",
        "template_version_id": "FPMS-FORMAT-LETTER-CUSTOMER-20260610-V1-006",
        "template_path": "templates/format_letters/format_letter_006.docx",
        "variables": COMMON_VARIABLES | {"deadline_text", "amount_lines_text"},
        "headers": (
            "序号",
            "贵方案号",
            "我方案号",
            "发明创造名称",
            "申请号",
            "申请日",
            "申请人名称",
            "文件名称",
            "登记期限",
            "登记金额",
        ),
        "cell_variables": {
            "登记期限": "deadline_text",
            "登记金额": "amount_lines_text",
        },
    },
    {
        "code": "FORMAT_LETTER_007",
        "customer_format_letter_name": "官文转发-国内客户-一通",
        "official_doc_name_pattern": "第一次审查意见通知书",
        "source_path": "docs/postdemo/文件样例及模版/常用邮件模板/审查意见或复审通知.doc",
        "source_sha256": "59322585ca96505fd1f38536ffa02e7afd3825df74c750e5256708b3869484cc",
        "template_version_id": "FPMS-FORMAT-LETTER-CUSTOMER-20260610-V1-007",
        "template_path": "templates/format_letters/format_letter_007.docx",
        "variables": COMMON_VARIABLES | {"deadline_text", "notice_variant_code"},
        "headers": (
            "序号",
            "贵方案号",
            "我方案号",
            "发明创造名称",
            "申请号",
            "申请日",
            "申请人名称",
            "文件名称",
            "答复期限",
        ),
        "cell_variables": {"答复期限": "deadline_text"},
    },
    {
        "code": "FORMAT_LETTER_008",
        "customer_format_letter_name": "官文转发-专利证书",
        "official_doc_name_pattern": "专利证书",
        "source_path": "docs/postdemo/文件样例及模版/常用邮件模板/专利证书.doc",
        "source_sha256": "5e9cf0638f02f98784c177e586f1127dc9fc0cc4e902596f61b368bb777c8be6",
        "template_version_id": "FPMS-FORMAT-LETTER-CUSTOMER-20260610-V1-008",
        "template_path": "templates/format_letters/format_letter_008.docx",
        "variables": COMMON_VARIABLES | {"inventor_names_text"},
        "headers": (
            "序号",
            "贵方案号",
            "我方案号",
            "发明创造名称",
            "申请号",
            "申请日",
            "发明人",
            "申请人名称",
            "文件名称",
        ),
        "cell_variables": {"发明人": "inventor_names_text"},
    },
)

COMMON_CELL_VARIABLES = {
    "贵方案号": "client_reference_no",
    "我方案号": "case_no",
    "发明创造名称": "invention_title",
    "申请号": "application_no",
    "申请日": "filing_date_text",
    "申请人名称": "applicant_names_text",
    "文件名称": "source_notice_name",
}


def _catalog_row_as_dict(row: object) -> dict[str, Any]:
    if is_dataclass(row):
        return asdict(row)
    if isinstance(row, Mapping):
        return dict(row)
    if hasattr(row, "_asdict"):
        return dict(row._asdict())
    pytest.fail("FORMAT_LETTER_MAPPING_CATALOG rows must be immutable named records")


def _catalog_projection(row: Mapping[str, Any]) -> dict[str, Any]:
    return {
        key: row[key]
        for key in (
            "code",
            "customer_format_letter_name",
            "official_doc_name_pattern",
            "source_path",
            "source_sha256",
            "template_version_id",
            "template_path",
        )
    }


def _template_path(row: Mapping[str, Any]) -> Path:
    return BACKEND_ROOT / "storage" / str(row["template_path"])


def _package_xml(path: Path) -> str:
    with zipfile.ZipFile(path) as package:
        return "\n".join(
            package.read(name).decode("utf-8", errors="ignore")
            for name in package.namelist()
            if name.endswith(".xml")
        )


def _document_text(document: Document) -> str:
    parts = [paragraph.text for paragraph in document.paragraphs]
    parts.extend(
        paragraph.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
    )
    return "\n".join(parts)


def _data_row_by_case_placeholder(document: Document):
    matching_rows = [
        row
        for row in document.tables[0].rows[1:]
        if "{{ case_no }}" in "\n".join(cell.text for cell in row.cells)
    ]
    assert len(matching_rows) == 1
    return matching_rows[0]


def _snapshot_format_letter_rows(
    db: Session,
) -> tuple[list[tuple[Any, ...]], list[tuple[Any, ...]]]:
    templates = db.execute(
        select(Template).where(Template.group == "FORMAT_LETTER").order_by(Template.id)
    ).scalars()
    mappings = db.execute(
        select(FormatLetterMapping)
        .where(FormatLetterMapping.format_letter_template_code.like("FORMAT_LETTER_%"))
        .order_by(FormatLetterMapping.id)
    ).scalars()
    return (
        [
            (row.id, row.name, row.group, row.language, row.file_path, row.enabled)
            for row in templates
        ],
        [
            (
                row.id,
                row.official_doc_name_pattern,
                row.format_letter_template_id,
                row.format_letter_template_code,
                row.output_name_rule,
                row.salutation_rule_code,
                row.contact_rule_code,
                row.enabled,
                row.remark,
            )
            for row in mappings
        ],
    )


def _copy_template_dataset(destination_base: Path) -> None:
    destination = destination_base / "storage" / "templates" / "format_letters"
    destination.mkdir(parents=True)
    for row in EXPECTED_ROWS:
        shutil.copy2(_template_path(row), destination / Path(row["template_path"]).name)


def test_catalog_and_customer_source_provenance_are_exact() -> None:
    catalog = seed_dev.FORMAT_LETTER_MAPPING_CATALOG
    assert isinstance(catalog, tuple)
    assert len(catalog) == 8
    normalized = tuple(_catalog_row_as_dict(row) for row in catalog)
    assert all(not isinstance(row, (dict, list, set)) for row in catalog)
    assert tuple(_catalog_projection(row) for row in normalized) == tuple(
        _catalog_projection(row) for row in EXPECTED_ROWS
    )

    for row in EXPECTED_ROWS:
        source = REPO_ROOT / row["source_path"]
        assert source.is_file()
        assert hashlib.sha256(source.read_bytes()).hexdigest() == row["source_sha256"]


def test_committed_templates_preserve_structure_and_exact_variables() -> None:
    for row in EXPECTED_ROWS:
        path = _template_path(row)
        assert path.is_file()
        variables = DocxTemplate(str(path)).get_undeclared_template_variables()
        assert variables == row["variables"], row["code"]

        package_xml = _package_xml(path)
        assert "MERGEFIELD" not in package_xml
        document = Document(path)
        text = _document_text(document)
        assert "请查收所附官方文件，并按期限完成后续事项" not in text
        assert "该专利申请因没有创造性而驳回" not in text
        assert re.search(r"发文日后\s*[1243]\s*个月", text) is None
        assert len(document.sections) == 1
        assert document.sections[0].page_width.twips == 16838
        assert document.sections[0].page_height.twips == 11906
        assert len(document.tables) == 1

        table = document.tables[0]
        assert tuple(cell.text.strip() for cell in table.rows[0].cells) == row["headers"]
        data_row = _data_row_by_case_placeholder(document)
        values_by_header = {
            header: cell.text for header, cell in zip(row["headers"], data_row.cells, strict=True)
        }
        expected_cells = COMMON_CELL_VARIABLES | row["cell_variables"]
        for header, variable in expected_cells.items():
            assert f"{{{{ {variable} }}}}" in values_by_header[header]

        assert "{{ salutation_text }}" in text
        assert text.count("{{ applicant_names_text }}") == 1


def test_templates_render_to_one_page_landscape_and_variant_007_has_one_data_row(
    tmp_path: Path,
) -> None:
    rendered_paths: list[Path] = []
    for row in EXPECTED_ROWS:
        variants = (
            ("OA_FIRST", "OA_SUBSEQUENT", "REEXAMINATION_NOTICE")
            if row["code"] == "FORMAT_LETTER_007"
            else ("OA_FIRST",)
        )
        for variant in variants:
            template = DocxTemplate(str(_template_path(row)))
            context = {**REPRESENTATIVE_CONTEXT, "notice_variant_code": variant}
            template.render(context)
            rendered_path = tmp_path / f"{row['code'].lower()}_{variant.lower()}.docx"
            template.save(rendered_path)
            rendered_paths.append(rendered_path)

            rendered = Document(rendered_path)
            rendered_text = _document_text(rendered)
            expected_values = {
                REPRESENTATIVE_CONTEXT[name]
                for name in row["variables"]
                if name != "notice_variant_code"
            }
            assert all(str(value) in rendered_text for value in expected_values)
            assert "{{" not in rendered_text
            assert "{%" not in rendered_text
            assert "MERGEFIELD" not in _package_xml(rendered_path)
            assert rendered.sections[0].page_width.twips == 16838
            assert rendered.sections[0].page_height.twips == 11906
            assert len(rendered.tables) == 1
            assert len(rendered.tables[0].rows) == 2

    soffice = shutil.which("soffice")
    pdfinfo = shutil.which("pdfinfo")
    assert soffice is not None, "soffice is required for frozen DOCX render verification"
    assert pdfinfo is not None, "pdfinfo is required for frozen PDF page verification"
    profile = tmp_path / "libreoffice-profile"
    profile.mkdir()
    result = subprocess.run(
        [
            soffice,
            "--headless",
            f"-env:UserInstallation={profile.resolve().as_uri()}",
            "--convert-to",
            "pdf",
            "--outdir",
            str(tmp_path),
            *(str(path) for path in rendered_paths),
        ],
        check=False,
        capture_output=True,
        text=True,
        timeout=60,
    )
    assert result.returncode == 0
    for rendered_path in rendered_paths:
        pdf_path = rendered_path.with_suffix(".pdf")
        assert pdf_path.is_file() and pdf_path.stat().st_size > 0
        info = subprocess.run(
            [pdfinfo, str(pdf_path)],
            check=True,
            capture_output=True,
            text=True,
            timeout=10,
        ).stdout
        assert re.search(r"^Pages:\s+1$", info, re.MULTILINE)
        size = re.search(r"^Page size:\s+([0-9.]+) x ([0-9.]+) pts", info, re.MULTILINE)
        assert size is not None
        width, height = map(float, size.groups())
        assert width > height
        assert width == pytest.approx(841.89, abs=1.0)
        assert height == pytest.approx(595.30, abs=1.0)


def test_seed_installs_exact_dataset_and_is_idempotent_without_commit(
    session_factory: sessionmaker,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    with session_factory() as db:
        monkeypatch.setattr(
            db,
            "commit",
            lambda: pytest.fail("seed_format_letter_mappings must not commit"),
        )
        monkeypatch.setattr(
            db,
            "rollback",
            lambda: pytest.fail("seed_format_letter_mappings must not roll back"),
        )
        assert seed_dev.seed_format_letter_mappings(db) == 16
        db.flush()

        templates = (
            db.execute(
                select(Template).where(Template.group == "FORMAT_LETTER").order_by(Template.name)
            )
            .scalars()
            .all()
        )
        mappings = (
            db.execute(
                select(FormatLetterMapping)
                .where(FormatLetterMapping.format_letter_template_code.like("FORMAT_LETTER_%"))
                .order_by(FormatLetterMapping.format_letter_template_code)
            )
            .scalars()
            .all()
        )
        assert len(templates) == len(mappings) == 8
        template_ids = {row.name: row.id for row in templates}
        mapping_ids = {row.format_letter_template_code: row.id for row in mappings}

        for expected, template, mapping in zip(EXPECTED_ROWS, templates, mappings, strict=True):
            assert (
                template.name,
                template.group,
                template.language,
                template.file_path,
                template.enabled,
            ) == (
                expected["code"],
                "FORMAT_LETTER",
                "zh-CN",
                expected["template_path"],
                True,
            )
            assert (
                mapping.official_doc_template_id,
                mapping.official_doc_template_code,
                mapping.official_doc_name_pattern,
                mapping.format_letter_template_id,
                mapping.format_letter_template_code,
                mapping.output_name_rule,
                mapping.salutation_rule_code,
                mapping.contact_rule_code,
                mapping.enabled,
            ) == (
                None,
                None,
                expected["official_doc_name_pattern"],
                template.id,
                expected["code"],
                OUTPUT_NAME_RULE,
                "PRIMARY_CONTACT_TITLE",
                "CLIENT_PRIMARY_CONTACT",
                True,
            )
            assert mapping.remark == json.dumps(
                {
                    "customer_format_letter_name": expected["customer_format_letter_name"],
                    "dataset_id": DATASET_ID,
                    "source_path": expected["source_path"],
                    "source_sha256": f"sha256:{expected['source_sha256']}",
                    "template_version_id": expected["template_version_id"],
                },
                ensure_ascii=False,
                sort_keys=True,
                separators=(",", ":"),
            )

        assert seed_dev.seed_format_letter_mappings(db) == 0
        db.flush()
        assert template_ids == {
            row.name: row.id
            for row in db.execute(
                select(Template).where(Template.group == "FORMAT_LETTER").order_by(Template.name)
            ).scalars()
        }
        assert mapping_ids == {
            row.format_letter_template_code: row.id
            for row in db.execute(
                select(FormatLetterMapping)
                .where(FormatLetterMapping.format_letter_template_code.like("FORMAT_LETTER_%"))
                .order_by(FormatLetterMapping.format_letter_template_code)
            ).scalars()
        }


def test_seed_updates_legacy_placeholder_rows_in_place(
    session_factory: sessionmaker,
) -> None:
    expected = EXPECTED_ROWS[0]
    template_id = str(uuid4())
    mapping_id = str(uuid4())
    with session_factory() as db:
        db.add(
            Template(
                id=template_id,
                name=expected["code"],
                group="FORMAT_LETTER",
                language="zh-CN",
                file_path="templates/format_letters/legacy-placeholder.docx",
                enabled=False,
            )
        )
        db.flush()
        db.add(
            FormatLetterMapping(
                id=mapping_id,
                official_doc_name_pattern="旧官文",
                format_letter_template_id=template_id,
                format_letter_template_code=expected["code"],
                output_name_rule="legacy.docx",
                enabled=False,
            )
        )
        db.flush()

        assert seed_dev.seed_format_letter_mappings(db) > 0
        db.flush()
        template = db.get(Template, template_id)
        mapping = db.get(FormatLetterMapping, mapping_id)
        assert template is not None and mapping is not None
        assert template.file_path == expected["template_path"]
        assert template.enabled is True
        assert mapping.official_doc_name_pattern == expected["official_doc_name_pattern"]
        assert mapping.enabled is True
        assert (
            len(
                db.execute(select(Template).where(Template.name == expected["code"]))
                .scalars()
                .all()
            )
            == 1
        )
        assert (
            len(
                db.execute(
                    select(FormatLetterMapping).where(
                        FormatLetterMapping.format_letter_template_code == expected["code"]
                    )
                )
                .scalars()
                .all()
            )
            == 1
        )


@pytest.mark.parametrize(
    ("mutation", "expected_prefix"),
    (
        ("missing", "FORMAT_LETTER_TEMPLATE_MISSING:FORMAT_LETTER_008"),
        ("corrupt", "FORMAT_LETTER_TEMPLATE_INVALID:FORMAT_LETTER_008"),
        (
            "wrong_variable",
            "FORMAT_LETTER_TEMPLATE_VARIABLES_MISMATCH:FORMAT_LETTER_008",
        ),
        (
            "placeholder",
            "FORMAT_LETTER_TEMPLATE_PLACEHOLDER_REMAINS:FORMAT_LETTER_008",
        ),
    ),
)
def test_template_validation_fails_before_seed_mutation_and_never_generates(
    mutation: str,
    expected_prefix: str,
    tmp_path: Path,
    session_factory: sessionmaker,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    temporary_base = tmp_path / "backend"
    _copy_template_dataset(temporary_base)
    target = temporary_base / "storage" / "templates" / "format_letters" / "format_letter_008.docx"
    if mutation == "missing":
        target.unlink()
    elif mutation == "corrupt":
        target.write_bytes(b"not-a-docx")
    elif mutation == "wrong_variable":
        document = Document(target)
        document.add_paragraph("{{ unsupported_variable }}")
        document.save(target)
    else:
        document = Document(target)
        document.add_paragraph("请查收所附官方文件，并按期限完成后续事项。")
        document.save(target)

    monkeypatch.setattr(seed_dev, "BASE_DIR", temporary_base)
    with session_factory() as db:
        before = _snapshot_format_letter_rows(db)
        with pytest.raises(RuntimeError, match=f"^{re.escape(expected_prefix)}"):
            seed_dev.seed_format_letter_mappings(db)
        db.flush()
        assert _snapshot_format_letter_rows(db) == before
    if mutation == "missing":
        assert not target.exists()


def test_ambiguous_database_carriers_fail_closed_before_mutation(
    session_factory: sessionmaker,
) -> None:
    code = EXPECTED_ROWS[0]["code"]
    with session_factory() as db:
        template = Template(
            id=str(uuid4()),
            name=code,
            group="FORMAT_LETTER",
            language="zh-CN",
            file_path=EXPECTED_ROWS[0]["template_path"],
            enabled=True,
        )
        db.add(template)
        db.flush()
        for _ in range(2):
            db.add(
                FormatLetterMapping(
                    id=str(uuid4()),
                    format_letter_template_id=template.id,
                    format_letter_template_code=code,
                    enabled=True,
                )
            )
        db.flush()
        before = _snapshot_format_letter_rows(db)
        with pytest.raises(
            RuntimeError,
            match=f"^FORMAT_LETTER_MAPPING_AMBIGUOUS:{code}",
        ):
            seed_dev.seed_format_letter_mappings(db)
        db.flush()
        assert _snapshot_format_letter_rows(db) == before


def test_late_catalog_mapping_ambiguity_preflight_leaves_entire_dataset_unchanged(
    session_factory: sessionmaker,
) -> None:
    code = EXPECTED_ROWS[-1]["code"]
    with session_factory() as db:
        for _ in range(2):
            db.add(
                FormatLetterMapping(
                    id=str(uuid4()),
                    format_letter_template_code=code,
                    enabled=True,
                )
            )
        db.flush()
        before = _snapshot_format_letter_rows(db)

        with pytest.raises(
            RuntimeError,
            match=f"^FORMAT_LETTER_MAPPING_AMBIGUOUS:{code}",
        ):
            seed_dev.seed_format_letter_mappings(db)

        db.flush()
        assert _snapshot_format_letter_rows(db) == before


@pytest.mark.parametrize("carrier_kind", ("duplicate_template", "unrelated_link"))
def test_ambiguous_or_unrelated_template_carrier_fails_closed(
    carrier_kind: str,
    session_factory: sessionmaker,
) -> None:
    code = EXPECTED_ROWS[0]["code"]
    with session_factory() as db:
        if carrier_kind == "duplicate_template":
            for _ in range(2):
                db.add(
                    Template(
                        id=str(uuid4()),
                        name=code,
                        group="FORMAT_LETTER",
                        language="zh-CN",
                        file_path=EXPECTED_ROWS[0]["template_path"],
                        enabled=True,
                    )
                )
        else:
            unrelated = Template(
                id=str(uuid4()),
                name="MANUAL_TEMPLATE",
                group="MANUAL",
                language="zh-CN",
                file_path="templates/manual.docx",
                enabled=True,
            )
            db.add(unrelated)
            db.flush()
            db.add(
                FormatLetterMapping(
                    id=str(uuid4()),
                    format_letter_template_id=unrelated.id,
                    format_letter_template_code=code,
                    enabled=True,
                )
            )
        db.flush()
        before = _snapshot_format_letter_rows(db)
        with pytest.raises(
            RuntimeError,
            match=f"^FORMAT_LETTER_TEMPLATE_ROW_AMBIGUOUS:{code}",
        ):
            seed_dev.seed_format_letter_mappings(db)
        db.flush()
        assert _snapshot_format_letter_rows(db) == before


def test_manual_unrelated_rows_are_not_deleted_or_normalized(
    session_factory: sessionmaker,
) -> None:
    manual_template_id = str(uuid4())
    manual_mapping_id = str(uuid4())
    with session_factory() as db:
        db.add(
            Template(
                id=manual_template_id,
                name="MANUAL_TEMPLATE",
                group="MANUAL",
                language="en-US",
                file_path="templates/manual.docx",
                enabled=False,
            )
        )
        db.add(
            FormatLetterMapping(
                id=manual_mapping_id,
                official_doc_name_pattern="手工映射",
                format_letter_template_code="MANUAL_TEMPLATE",
                output_name_rule="manual.docx",
                enabled=False,
            )
        )
        db.flush()
        assert seed_dev.seed_format_letter_mappings(db) == 16
        db.flush()
        manual_template = db.get(Template, manual_template_id)
        manual_mapping = db.get(FormatLetterMapping, manual_mapping_id)
        assert manual_template is not None and manual_mapping is not None
        assert (
            manual_template.group,
            manual_template.language,
            manual_template.file_path,
            manual_template.enabled,
        ) == ("MANUAL", "en-US", "templates/manual.docx", False)
        assert (
            manual_mapping.official_doc_name_pattern,
            manual_mapping.output_name_rule,
            manual_mapping.enabled,
        ) == ("手工映射", "manual.docx", False)
