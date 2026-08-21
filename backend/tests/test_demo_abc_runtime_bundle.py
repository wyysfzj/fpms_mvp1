from __future__ import annotations

import hashlib
import json
import tomllib
import zipfile
from datetime import date
from pathlib import Path

import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from pypdf import PdfWriter
from pypdf.generic import (
    ArrayObject,
    DecodedStreamObject,
    DictionaryObject,
    NameObject,
    NumberObject,
    TextStringObject,
)

try:
    from app.core import demo_bundle
    from app.core.demo_bundle import DemoBundleError, load_demo_bundle
except ImportError:
    demo_bundle = None

    class DemoBundleError(RuntimeError):
        pass

    def load_demo_bundle(*_args, **_kwargs):
        pytest.fail("demo bundle preflight is not implemented")


REPO_ROOT = Path(__file__).resolve().parents[2]
CONTRACT_REF = "docs/superpowers/specs/2026-08-15-fpms-local-demo-abc-design.md"
INTEGRATED_CONTRACT_REF = "docs/superpowers/specs/2026-08-21-fpms-integrated-demo-a-design.md"
INTEGRATED_DECISION_REF = (
    "docs/product/v8/customer-decisions/"
    "2026-08-21-integrated-demo-a-written-spec-acceptance.txt"
)
EVIDENCE_ROLES = [
    "FILING_FINAL_SUBMISSION",
    "FILING_RECEIPT",
    "ACCEPTANCE_NOTICE",
    "PRELIMINARY_EXAMINATION_SOURCE",
    "PUBLICATION_NOTICE",
    "SUBSTANTIVE_EXAMINATION_SOURCE",
    "OA_NOTICE",
    "OA_RECEIPT",
]
INTEGRATED_EVIDENCE_ROLES = [
    "FILING_FINAL_SUBMISSION",
    "FILING_RECEIPT",
    "ACCEPTANCE_NOTICE",
    "PRELIMINARY_EXAMINATION_SOURCE",
    "PUBLICATION_NOTICE",
    "SUBSTANTIVE_EXAMINATION_SOURCE",
    "OA_NOTICE_1",
    "OA_RECEIPT_1",
    "OA_NOTICE_2",
    "OA_RECEIPT_2",
    "GRANT_NOTICE_ORIGINAL",
    "GRANT_NOTICE_REPLACEMENT",
]


def _sha256(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def _write_docx(
    path: Path,
    *,
    marker: bool = True,
    hidden_marker: bool = False,
    style_hidden_marker: bool = False,
    default_hidden_marker: str | None = None,
    deleted_marker: bool = False,
    external: bool = False,
) -> None:
    marker_text = "DEMO_ONLY / 仅用于本地虚构演示" if marker else "普通模板"
    document = Document()
    marker_run = document.add_paragraph().add_run(marker_text)
    marker_run.font.hidden = hidden_marker
    if style_hidden_marker:
        hidden_base = document.styles.add_style("HiddenMarkerBase", WD_STYLE_TYPE.CHARACTER)
        hidden_base.font.hidden = True
        hidden_child = document.styles.add_style("HiddenMarkerChild", WD_STYLE_TYPE.CHARACTER)
        hidden_child.base_style = hidden_base
        marker_run.style = hidden_child
    if default_hidden_marker == "paragraph":
        document.styles["Normal"].font.hidden = True
    elif default_hidden_marker == "character":
        document.styles["Default Paragraph Font"].font.hidden = True
    if deleted_marker:
        paragraph = marker_run._r.getparent()
        deleted = OxmlElement("w:del")
        paragraph.remove(marker_run._r)
        deleted.append(marker_run._r)
        paragraph.append(deleted)
    document.add_paragraph("案号 {{ case_no }} / 客户 {{ client_name }}")
    document.save(path)
    if external:
        with zipfile.ZipFile(path, "a", compression=zipfile.ZIP_DEFLATED) as archive:
            archive.writestr(
                "word/_rels/demo-external.rels",
                '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
                '<Relationship Id="rId999" Target="https://example.invalid" '
                'TargetMode="External" /></Relationships>',
            )


def _write_pdf(
    path: Path,
    *,
    marker: str = "bilingual",
    invisible: bool = False,
    font_size: int = 12,
    horizontal_scale: int = 100,
    mixed_horizontal_scale: str | None = None,
    unsafe_operator: str | None = None,
    unique_text: str | None = None,
) -> None:
    writer = PdfWriter()
    if unique_text is not None:
        writer.add_metadata({"/Title": unique_text})
    page = writer.add_blank_page(width=612, height=792)
    descendant_font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/CIDFontType0"),
            NameObject("/BaseFont"): NameObject("/STSong-Light"),
            NameObject("/CIDSystemInfo"): DictionaryObject(
                {
                    NameObject("/Registry"): TextStringObject("Adobe"),
                    NameObject("/Ordering"): TextStringObject("GB1"),
                    NameObject("/Supplement"): NumberObject(4),
                }
            ),
        }
    )
    descendant_ref = writer._add_object(descendant_font)
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type0"),
            NameObject("/BaseFont"): NameObject("/STSong-Light"),
            NameObject("/Encoding"): NameObject("/UniGB-UCS2-H"),
            NameObject("/DescendantFonts"): ArrayObject([descendant_ref]),
        }
    )
    font_ref = writer._add_object(font)
    page[NameObject("/Resources")] = DictionaryObject(
        {NameObject("/Font"): DictionaryObject({NameObject("/F1"): font_ref})}
    )
    stream = DecodedStreamObject()
    visible = {
        "bilingual": "FICTIONAL_DEMO_EVIDENCE / 仅用于本地虚构演示",
        "english": "FICTIONAL_DEMO_EVIDENCE / LOCAL FICTIONAL DEMO",
        "missing": "ordinary",
    }[marker]
    encoded = visible.encode("utf-16-be").hex().upper()
    render_mode = "3 Tr " if invisible else ""
    if unsafe_operator == "quote":
        content = (
            f"BT /F1 {font_size} Tf 100 Tz 1000 TL 72 720 Td <{encoded}> ' ET"
        )
    elif unsafe_operator == "double-quote":
        content = (
            f'BT /F1 {font_size} Tf 100 Tz 1000 TL 72 720 Td 0 0 <{encoded}> " ET'
        )
    elif unsafe_operator == "tj-array":
        content = (
            f"BT /F1 {font_size} Tf 100 Tz 72 720 Td [100000 <{encoded}>] TJ ET"
        )
    elif mixed_horizontal_scale is None:
        content = (
            f"BT /F1 {font_size} Tf {horizontal_scale} Tz "
            f"{render_mode}72 720 Td <{encoded}> Tj ET"
        )
    else:
        ordinary = "ordinary".encode("utf-16-be").hex().upper()
        if mixed_horizontal_scale == "hidden-marker-first":
            content = (
                f"BT /F1 {font_size} Tf 0 Tz 72 720 Td <{encoded}> Tj "
                f"100 Tz <{ordinary}> Tj ET"
            )
        elif mixed_horizontal_scale == "visible-marker-first":
            content = (
                f"BT /F1 {font_size} Tf 100 Tz 72 720 Td <{encoded}> Tj "
                f"0 Tz <{ordinary}> Tj ET"
            )
        elif mixed_horizontal_scale == "hidden-marker-then-state":
            content = (
                f"BT /F1 {font_size} Tf 0 Tz 72 720 Td <{encoded}> Tj 100 Tz ET"
            )
        else:
            content = (
                f"BT /F1 {font_size} Tf 100 Tz 72 720 Td <{encoded}> Tj 0 Tz ET"
            )
    stream.set_data(content.encode())
    page[NameObject("/Contents")] = writer._add_object(stream)
    with path.open("wb") as output:
        writer.write(output)


def _metadata_for(role: str) -> dict[str, object]:
    metadata: dict[str, object] = {
        "effective_at": None,
        "received_at": None,
        "receipt_kind": None,
        "official_due_date": None,
        "official_due_date_source": None,
        "official_due_date_status": None,
        "oa_sequence": None,
        "source_template_code": None,
    }
    if role in {"FILING_RECEIPT", "OA_RECEIPT"}:
        metadata["received_at"] = "2026-08-16T10:00:00"
        metadata["receipt_kind"] = "RECEIPT_PDF"
    else:
        metadata["effective_at"] = "2026-08-16T09:00:00"
    if role == "OA_NOTICE":
        metadata.update(
            official_due_date="2026-09-16",
            official_due_date_source="MANUAL_OFFICIAL_NOTICE",
            official_due_date_status="CONFIRMED",
            source_template_code="DEMO_OA_NOTICE_1",
            oa_sequence=1,
        )
    return metadata


def _integrated_metadata_for(role: str, ordinal: int) -> dict[str, object]:
    metadata: dict[str, object] = {
        "effective_at": None,
        "received_at": None,
        "receipt_kind": None,
        "official_due_date": None,
        "official_due_date_source": None,
        "official_due_date_status": None,
        "oa_sequence": None,
        "source_template_code": None,
        "supersedes_role": None,
    }
    if role in {"FILING_RECEIPT", "OA_RECEIPT_1", "OA_RECEIPT_2"}:
        metadata["received_at"] = f"2026-08-{ordinal + 1:02d}T10:00:00"
        metadata["receipt_kind"] = "RECEIPT_PDF"
        return metadata

    metadata["effective_at"] = f"2026-08-{ordinal + 1:02d}T09:00:00"
    if role in {"OA_NOTICE_1", "OA_NOTICE_2"}:
        sequence = 1 if role.endswith("_1") else 2
        metadata.update(
            official_due_date=f"2026-{9 if sequence == 1 else 10:02d}-{21 + sequence:02d}",
            official_due_date_source="MANUAL_OFFICIAL_NOTICE",
            official_due_date_status="CONFIRMED",
            oa_sequence=sequence,
            source_template_code=f"DEMO_OA_NOTICE_{sequence}",
        )
    elif role in {"GRANT_NOTICE_ORIGINAL", "GRANT_NOTICE_REPLACEMENT"}:
        replacement = role == "GRANT_NOTICE_REPLACEMENT"
        metadata.update(
            official_due_date="2026-11-24" if replacement else "2026-11-23",
            official_due_date_source="IMPORTED_OFFICIAL_NOTICE",
            official_due_date_status="CONFIRMED",
            source_template_code=(
                "DEMO_GRANT_NOTICE_2" if replacement else "DEMO_GRANT_NOTICE_1"
            ),
            supersedes_role=("GRANT_NOTICE_ORIGINAL" if replacement else None),
        )
    return metadata


def _write_manifest(bundle_root: Path, manifest: dict[str, object]) -> str:
    raw = (json.dumps(manifest, ensure_ascii=False, separators=(",", ":")) + "\n").encode()
    (bundle_root / "manifest.json").write_bytes(raw)
    manifest_digest = _sha256(raw)
    decision_ref = manifest["authority"]["decision_ref"]
    file_rows = [manifest["templates"][0], *manifest["evidence"]]
    authority = {
        "schema_version": "fpms.demo-bundle-authority/v1",
        "status": "APPROVED",
        "authority_classification": manifest["authority_classification"],
        "approved_by": "synthetic-test-fixture-generator",
        "approved_at": "2026-08-16T12:00:00+08:00",
        "decision_ref": decision_ref,
        "decision_version": manifest["authority"]["decision_version"],
        "decision_sha256": _sha256((REPO_ROOT / decision_ref).read_bytes()),
        "bundle_id": manifest["bundle_id"],
        "bundle_version": manifest["bundle_version"],
        "manifest_sha256": manifest_digest,
        "source_digests": [
            {
                "kind": "PROVENANCE",
                "ref": manifest["provenance"]["source_ref"],
                "version": manifest["provenance"]["source_version"],
                "sha256": manifest["provenance"]["source_sha256"],
            },
            {
                "kind": "SERVICE_RATE",
                "ref": manifest["rates"][0]["source_ref"],
                "version": manifest["rates"][0]["source_version"],
                "sha256": manifest["rates"][0]["source_sha256"],
            },
        ],
        "file_digests": sorted(
            ({"path": row["path"], "sha256": row["sha256"]} for row in file_rows),
            key=lambda row: row["path"],
        ),
    }
    authority_raw = (
        json.dumps(authority, ensure_ascii=False, separators=(",", ":")) + "\n"
    ).encode()
    (bundle_root / "authority.json").write_bytes(authority_raw)
    return manifest_digest


def _authority_digest(bundle_root: Path) -> str:
    return _sha256((bundle_root / "authority.json").read_bytes())


def _load_bundle(bundle_root: Path, manifest_digest: str):
    return load_demo_bundle(
        bundle_root,
        expected_manifest_sha256=manifest_digest,
        expected_authority_sha256=_authority_digest(bundle_root),
        expected_authority_classification="SYNTHETIC_TEST_ONLY",
        repo_root=REPO_ROOT,
    )


def _valid_bundle(tmp_path: Path) -> tuple[Path, dict[str, object], str]:
    bundle_root = tmp_path / "bundle"
    templates = bundle_root / "templates"
    evidence = bundle_root / "evidence"
    templates.mkdir(parents=True)
    evidence.mkdir()

    template_path = templates / "demo-letter.docx"
    _write_docx(template_path)
    evidence_rows: list[dict[str, object]] = []
    for role in EVIDENCE_ROLES:
        evidence_path = evidence / f"{role.lower()}.pdf"
        _write_pdf(evidence_path)
        evidence_bytes = evidence_path.read_bytes()
        evidence_rows.append(
            {
                "role": role,
                "title_zh_cn": f"虚构演示证据-{role}",
                "classification": "FICTIONAL_DEMO_EVIDENCE",
                "path": f"evidence/{evidence_path.name}",
                "media_type": "application/pdf",
                "size_bytes": len(evidence_bytes),
                "sha256": _sha256(evidence_bytes),
                "metadata": _metadata_for(role),
            }
        )

    contract_bytes = (REPO_ROOT / CONTRACT_REF).read_bytes()
    manifest: dict[str, object] = {
        "schema_version": "fpms.demo-input-bundle/v1",
        "bundle_id": "fpms-local-abc",
        "bundle_version": "2026.08.16",
        "classification": "DEMO_ONLY",
        "authority_classification": "SYNTHETIC_TEST_ONLY",
        "purpose": "LOCAL_ABC_E2E",
        "valid_from": "2026-08-16",
        "valid_until": "2026-08-31",
        "authority": {
            "decision_ref": "docs/product/v8/customer-decisions/2026-08-15-local-demo-abc.txt",
            "decision_version": "DEC-LOCAL-DEMO-ABC-20260815",
        },
        "provenance": {
            "label_zh_cn": "本地虚构演示输入",
            "source_ref": "synthetic-test-only-input",
            "source_version": "2026.08.16",
            "source_sha256": "a" * 64,
        },
        "contract": {"ref": CONTRACT_REF, "sha256": _sha256(contract_bytes)},
        "capabilities": [
            "FICTIONAL_LIFECYCLE_EVIDENCE",
            "INTERNAL_TEMPLATE_PREVIEW",
            "SERVICE_PRICE_TO_OBLIGATION",
        ],
        "templates": [
            {
                "consumer": "DOCUMENT_RENDER",
                "template_code": "DEMO_INTERNAL_LETTER_1",
                "group": "INTERNAL_DEMO",
                "language": "zh-CN",
                "path": "templates/demo-letter.docx",
                "media_type": (
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                ),
                "size_bytes": template_path.stat().st_size,
                "sha256": _sha256(template_path.read_bytes()),
                "required_variables": ["case_no", "client_name"],
            }
        ],
        "evidence": evidence_rows,
        "rates": [
            {
                "domain": "SERVICE_DEMO_PRICE",
                "item_code": "DEMO_SERVICE_1",
                "name_zh_cn": "演示服务费",
                "currency": "CNY",
                "calc_mode": "FIXED",
                "amount": "1200.00",
                "source_ref": "synthetic-test-only-rate",
                "source_version": "2026.08.16",
                "source_sha256": "b" * 64,
                "disclaimer_zh_cn": "仅用于本地虚构演示，不是正式报价或官方费用。",
            }
        ],
    }
    return bundle_root, manifest, _write_manifest(bundle_root, manifest)


def _valid_integrated_bundle(tmp_path: Path) -> tuple[Path, dict[str, object], str]:
    bundle_root = tmp_path / "bundle"
    templates = bundle_root / "templates"
    evidence = bundle_root / "evidence"
    templates.mkdir(parents=True)
    evidence.mkdir()

    template_path = templates / "integrated-demo-letter.docx"
    _write_docx(template_path)
    evidence_rows: list[dict[str, object]] = []
    for ordinal, role in enumerate(INTEGRATED_EVIDENCE_ROLES):
        evidence_path = evidence / f"{ordinal + 1:02d}-{role.lower()}.pdf"
        _write_pdf(evidence_path, unique_text=f"integrated-a-{ordinal + 1:02d}-{role}")
        evidence_rows.append(
            {
                "role": role,
                "title_zh_cn": f"虚构集成演示证据-{ordinal + 1:02d}",
                "classification": "FICTIONAL_DEMO_EVIDENCE",
                "path": f"evidence/{evidence_path.name}",
                "media_type": "application/pdf",
                "size_bytes": evidence_path.stat().st_size,
                "sha256": _sha256(evidence_path.read_bytes()),
                "metadata": _integrated_metadata_for(role, ordinal),
            }
        )

    contract_bytes = (REPO_ROOT / INTEGRATED_CONTRACT_REF).read_bytes()
    manifest: dict[str, object] = {
        "schema_version": "fpms.demo-input-bundle/integrated-a-v1",
        "bundle_id": "fpms-integrated-a",
        "bundle_version": "2026.08.21",
        "classification": "DEMO_ONLY",
        "authority_classification": "SYNTHETIC_TEST_ONLY",
        "purpose": "LOCAL_INTEGRATED_A_E2E",
        "valid_from": "2026-08-21",
        "valid_until": "2026-09-30",
        "authority": {
            "decision_ref": INTEGRATED_DECISION_REF,
            "decision_version": "DEC-INTEGRATED-DEMO-A-20260821",
        },
        "provenance": {
            "label_zh_cn": "本地虚构集成演示输入",
            "source_ref": "synthetic-integrated-a-input",
            "source_version": "2026.08.21",
            "source_sha256": "c" * 64,
        },
        "contract": {
            "ref": INTEGRATED_CONTRACT_REF,
            "sha256": _sha256(contract_bytes),
        },
        "capabilities": [
            "FICTIONAL_LIFECYCLE_EVIDENCE",
            "INTERNAL_TEMPLATE_PREVIEW",
            "SERVICE_PRICE_TO_OBLIGATION",
        ],
        "templates": [
            {
                "consumer": "DOCUMENT_RENDER",
                "template_code": "DEMO_INTEGRATED_LETTER_1",
                "group": "INTERNAL_DEMO",
                "language": "zh-CN",
                "path": "templates/integrated-demo-letter.docx",
                "media_type": (
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                ),
                "size_bytes": template_path.stat().st_size,
                "sha256": _sha256(template_path.read_bytes()),
                "required_variables": ["case_no", "client_name"],
            }
        ],
        "evidence": evidence_rows,
        "rates": [
            {
                "domain": "SERVICE_DEMO_PRICE",
                "item_code": "DEMO_INTEGRATED_SERVICE_1",
                "name_zh_cn": "集成演示服务费",
                "currency": "CNY",
                "calc_mode": "FIXED",
                "amount": "1200.00",
                "source_ref": "synthetic-integrated-a-rate",
                "source_version": "2026.08.21",
                "source_sha256": "d" * 64,
                "disclaimer_zh_cn": "仅用于本地虚构集成演示，不是正式报价或官方费用。",
            }
        ],
    }
    return bundle_root, manifest, _write_manifest(bundle_root, manifest)


def test_valid_bundle_returns_immutable_snapshot(tmp_path: Path, monkeypatch: pytest.MonkeyPatch):
    dependencies = tomllib.loads((REPO_ROOT / "backend/pyproject.toml").read_text())["project"][
        "dependencies"
    ]
    assert "pypdf>=6.0,<7" in dependencies
    root, _manifest, digest = _valid_bundle(tmp_path)
    monkeypatch.setattr(demo_bundle, "_current_demo_date", lambda: date(2026, 8, 16))

    snapshot = _load_bundle(root, digest)

    assert snapshot.bundle_id == "fpms-local-abc"
    assert snapshot.manifest_sha256 == digest
    assert snapshot.authority_sha256 == _authority_digest(root)
    assert snapshot.authority_classification == "SYNTHETIC_TEST_ONLY"
    assert snapshot.customer_activation_eligible is False
    assert snapshot.approved_by == "synthetic-test-fixture-generator"
    assert snapshot.approved_at == "2026-08-16T12:00:00+08:00"
    assert snapshot.service_rate.amount == "1200.00"
    assert snapshot.evidence_roles == tuple(EVIDENCE_ROLES)


@pytest.mark.parametrize(
    ("mutation", "message"),
    [
        (lambda manifest: manifest.update(extra=True), "unknown keys"),
        (
            lambda manifest: manifest["evidence"][6]["metadata"].update(
                source_template_code="OTHER_OA"
            ),
            "OA semantic",
        ),
        (lambda manifest: manifest.update(valid_until="2026-08-15"), "validity"),
    ],
)
def test_manifest_contract_mismatch_fails_closed(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    mutation,
    message: str,
):
    root, manifest, _digest = _valid_bundle(tmp_path)
    mutation(manifest)
    digest = _write_manifest(root, manifest)
    monkeypatch.setattr(demo_bundle, "_current_demo_date", lambda: date(2026, 8, 16))

    with pytest.raises(DemoBundleError, match=message):
        _load_bundle(root, digest)


def test_external_manifest_digest_is_checked_before_parsing(tmp_path: Path):
    root, _manifest, _digest = _valid_bundle(tmp_path)

    with pytest.raises(DemoBundleError, match="manifest digest"):
        load_demo_bundle(
            root,
            expected_manifest_sha256="0" * 64,
            expected_authority_sha256=_authority_digest(root),
            expected_authority_classification="SYNTHETIC_TEST_ONLY",
            repo_root=REPO_ROOT,
        )


def test_file_hash_extra_file_and_marker_fail_closed(tmp_path: Path, monkeypatch):
    root, manifest, digest = _valid_bundle(tmp_path)
    monkeypatch.setattr(demo_bundle, "_current_demo_date", lambda: date(2026, 8, 16))
    evidence_path = root / manifest["evidence"][0]["path"]
    evidence_path.write_bytes(b"tampered")

    with pytest.raises(DemoBundleError, match="size or hash"):
        _load_bundle(root, digest)

    root, _manifest, digest = _valid_bundle(tmp_path / "extra")
    (root / "unexpected.txt").write_text("unexpected")
    with pytest.raises(DemoBundleError, match="file set"):
        _load_bundle(root, digest)

    root, manifest, _digest = _valid_bundle(tmp_path / "marker")
    template_path = root / manifest["templates"][0]["path"]
    _write_docx(template_path, marker=False)
    manifest["templates"][0]["size_bytes"] = template_path.stat().st_size
    manifest["templates"][0]["sha256"] = _sha256(template_path.read_bytes())
    digest = _write_manifest(root, manifest)
    with pytest.raises(DemoBundleError, match="demo marker"):
        _load_bundle(root, digest)

    root, manifest, _digest = _valid_bundle(tmp_path / "hidden-marker")
    template_path = root / manifest["templates"][0]["path"]
    _write_docx(template_path, hidden_marker=True)
    manifest["templates"][0]["size_bytes"] = template_path.stat().st_size
    manifest["templates"][0]["sha256"] = _sha256(template_path.read_bytes())
    digest = _write_manifest(root, manifest)
    with pytest.raises(DemoBundleError, match="visible demo marker"):
        _load_bundle(root, digest)

    for default_kind in ("paragraph", "character"):
        root, manifest, _digest = _valid_bundle(
            tmp_path / f"default-hidden-{default_kind}"
        )
        template_path = root / manifest["templates"][0]["path"]
        _write_docx(template_path, default_hidden_marker=default_kind)
        manifest["templates"][0]["size_bytes"] = template_path.stat().st_size
        manifest["templates"][0]["sha256"] = _sha256(template_path.read_bytes())
        digest = _write_manifest(root, manifest)
        with pytest.raises(DemoBundleError, match="visible demo marker"):
            _load_bundle(root, digest)

    root, manifest, _digest = _valid_bundle(tmp_path / "style-hidden-marker")
    template_path = root / manifest["templates"][0]["path"]
    _write_docx(template_path, style_hidden_marker=True)
    manifest["templates"][0]["size_bytes"] = template_path.stat().st_size
    manifest["templates"][0]["sha256"] = _sha256(template_path.read_bytes())
    digest = _write_manifest(root, manifest)
    with pytest.raises(DemoBundleError, match="visible demo marker"):
        _load_bundle(root, digest)

    root, manifest, _digest = _valid_bundle(tmp_path / "deleted-marker")
    template_path = root / manifest["templates"][0]["path"]
    _write_docx(template_path, deleted_marker=True)
    manifest["templates"][0]["size_bytes"] = template_path.stat().st_size
    manifest["templates"][0]["sha256"] = _sha256(template_path.read_bytes())
    digest = _write_manifest(root, manifest)
    with pytest.raises(DemoBundleError, match="visible demo marker"):
        _load_bundle(root, digest)

    root, manifest, _digest = _valid_bundle(tmp_path / "english-pdf")
    evidence_path = root / manifest["evidence"][0]["path"]
    _write_pdf(evidence_path, marker="english")
    manifest["evidence"][0]["size_bytes"] = evidence_path.stat().st_size
    manifest["evidence"][0]["sha256"] = _sha256(evidence_path.read_bytes())
    digest = _write_manifest(root, manifest)
    with pytest.raises(DemoBundleError, match="bilingual"):
        _load_bundle(root, digest)

    root, manifest, _digest = _valid_bundle(tmp_path / "invisible-pdf")
    evidence_path = root / manifest["evidence"][0]["path"]
    _write_pdf(evidence_path, invisible=True)
    manifest["evidence"][0]["size_bytes"] = evidence_path.stat().st_size
    manifest["evidence"][0]["sha256"] = _sha256(evidence_path.read_bytes())
    digest = _write_manifest(root, manifest)
    with pytest.raises(DemoBundleError, match="visible bilingual"):
        _load_bundle(root, digest)

    for order in ("hidden-marker-first", "visible-marker-first"):
        root, manifest, _digest = _valid_bundle(tmp_path / f"mixed-tz-{order}")
        evidence_path = root / manifest["evidence"][0]["path"]
        _write_pdf(evidence_path, mixed_horizontal_scale=order)
        manifest["evidence"][0]["size_bytes"] = evidence_path.stat().st_size
        manifest["evidence"][0]["sha256"] = _sha256(evidence_path.read_bytes())
        digest = _write_manifest(root, manifest)
        with pytest.raises(DemoBundleError, match="visible bilingual"):
            _load_bundle(root, digest)

    root, manifest, _digest = _valid_bundle(tmp_path / "hidden-marker-then-state")
    evidence_path = root / manifest["evidence"][0]["path"]
    _write_pdf(evidence_path, mixed_horizontal_scale="hidden-marker-then-state")
    manifest["evidence"][0]["size_bytes"] = evidence_path.stat().st_size
    manifest["evidence"][0]["sha256"] = _sha256(evidence_path.read_bytes())
    digest = _write_manifest(root, manifest)
    with pytest.raises(DemoBundleError, match="visible bilingual"):
        _load_bundle(root, digest)

    root, manifest, _digest = _valid_bundle(tmp_path / "visible-marker-then-state")
    evidence_path = root / manifest["evidence"][0]["path"]
    _write_pdf(evidence_path, mixed_horizontal_scale="visible-marker-then-state")
    manifest["evidence"][0]["size_bytes"] = evidence_path.stat().st_size
    manifest["evidence"][0]["sha256"] = _sha256(evidence_path.read_bytes())
    digest = _write_manifest(root, manifest)
    _load_bundle(root, digest)

    for operator in ("quote", "double-quote", "tj-array"):
        root, manifest, _digest = _valid_bundle(tmp_path / f"unsafe-{operator}")
        evidence_path = root / manifest["evidence"][0]["path"]
        _write_pdf(evidence_path, unsafe_operator=operator)
        manifest["evidence"][0]["size_bytes"] = evidence_path.stat().st_size
        manifest["evidence"][0]["sha256"] = _sha256(evidence_path.read_bytes())
        digest = _write_manifest(root, manifest)
        with pytest.raises(DemoBundleError, match="visible bilingual"):
            _load_bundle(root, digest)

    root, manifest, _digest = _valid_bundle(tmp_path / "zero-horizontal-scale-pdf")
    evidence_path = root / manifest["evidence"][0]["path"]
    _write_pdf(evidence_path, horizontal_scale=0)
    manifest["evidence"][0]["size_bytes"] = evidence_path.stat().st_size
    manifest["evidence"][0]["sha256"] = _sha256(evidence_path.read_bytes())
    digest = _write_manifest(root, manifest)
    with pytest.raises(DemoBundleError, match="visible bilingual"):
        _load_bundle(root, digest)

    root, manifest, _digest = _valid_bundle(tmp_path / "zero-font-pdf")
    evidence_path = root / manifest["evidence"][0]["path"]
    _write_pdf(evidence_path, font_size=0)
    manifest["evidence"][0]["size_bytes"] = evidence_path.stat().st_size
    manifest["evidence"][0]["sha256"] = _sha256(evidence_path.read_bytes())
    digest = _write_manifest(root, manifest)
    with pytest.raises(DemoBundleError, match="visible bilingual"):
        _load_bundle(root, digest)


def test_authority_classification_is_cross_bound_and_expected(tmp_path: Path, monkeypatch):
    root, manifest, _digest = _valid_bundle(tmp_path)
    monkeypatch.setattr(demo_bundle, "_current_demo_date", lambda: date(2026, 8, 16))
    manifest["authority_classification"] = "CUSTOMER_AUTHORIZED"
    digest = _write_manifest(root, manifest)

    with pytest.raises(DemoBundleError, match="authority classification"):
        _load_bundle(root, digest)


@pytest.mark.parametrize("oa_sequence", [True, 1.0, "1"])
def test_oa_sequence_requires_exact_integer_one(tmp_path: Path, monkeypatch, oa_sequence):
    root, manifest, _digest = _valid_bundle(tmp_path)
    monkeypatch.setattr(demo_bundle, "_current_demo_date", lambda: date(2026, 8, 16))
    manifest["evidence"][6]["metadata"]["oa_sequence"] = oa_sequence
    digest = _write_manifest(root, manifest)

    with pytest.raises(DemoBundleError, match="oa_sequence"):
        _load_bundle(root, digest)


def test_docx_external_relationship_fails_closed(tmp_path: Path, monkeypatch):
    root, manifest, _digest = _valid_bundle(tmp_path)
    monkeypatch.setattr(demo_bundle, "_current_demo_date", lambda: date(2026, 8, 16))
    template_path = root / manifest["templates"][0]["path"]
    _write_docx(template_path, external=True)
    manifest["templates"][0]["size_bytes"] = template_path.stat().st_size
    manifest["templates"][0]["sha256"] = _sha256(template_path.read_bytes())
    digest = _write_manifest(root, manifest)

    with pytest.raises(DemoBundleError, match="external relationship"):
        _load_bundle(root, digest)


def test_docx_placeholder_drift_and_pseudo_pdf_fail_closed(tmp_path: Path, monkeypatch):
    root, manifest, _digest = _valid_bundle(tmp_path)
    monkeypatch.setattr(demo_bundle, "_current_demo_date", lambda: date(2026, 8, 16))
    manifest["templates"][0]["required_variables"] = ["case_no", "missing_value"]
    digest = _write_manifest(root, manifest)
    with pytest.raises(DemoBundleError, match="placeholder"):
        _load_bundle(root, digest)

    root, manifest, _digest = _valid_bundle(tmp_path / "pdf")
    evidence_path = root / manifest["evidence"][0]["path"]
    evidence_path.write_bytes(
        b"%PDF-1.4\n% FICTIONAL_DEMO_EVIDENCE / marker only in comment\n%%EOF\n"
    )
    manifest["evidence"][0]["size_bytes"] = evidence_path.stat().st_size
    manifest["evidence"][0]["sha256"] = _sha256(evidence_path.read_bytes())
    digest = _write_manifest(root, manifest)
    with pytest.raises(DemoBundleError, match="PDF"):
        _load_bundle(root, digest)


def test_bundle_root_symlink_fails_closed(tmp_path: Path, monkeypatch):
    root, _manifest, digest = _valid_bundle(tmp_path / "source")
    monkeypatch.setattr(demo_bundle, "_current_demo_date", lambda: date(2026, 8, 16))
    alias = tmp_path / "bundle-link"
    alias.symlink_to(root, target_is_directory=True)

    with pytest.raises(DemoBundleError, match="root"):
        load_demo_bundle(
            alias,
            expected_manifest_sha256=digest,
            expected_authority_sha256=_authority_digest(root),
            expected_authority_classification="SYNTHETIC_TEST_ONLY",
            repo_root=REPO_ROOT,
        )


def test_bundle_root_inside_product_storage_fails_closed(tmp_path: Path, monkeypatch):
    storage_root = tmp_path / "product-storage"
    root, _manifest, digest = _valid_bundle(storage_root / "input")
    monkeypatch.setattr(demo_bundle, "_current_demo_date", lambda: date(2026, 8, 16))

    with pytest.raises(DemoBundleError, match="product and run storage"):
        load_demo_bundle(
            root,
            expected_manifest_sha256=digest,
            expected_authority_sha256=_authority_digest(root),
            expected_authority_classification="SYNTHETIC_TEST_ONLY",
            repo_root=REPO_ROOT,
            forbidden_roots=(storage_root,),
        )


def test_authority_record_is_independently_pinned_and_exact(tmp_path: Path, monkeypatch):
    root, _manifest, digest = _valid_bundle(tmp_path)
    monkeypatch.setattr(demo_bundle, "_current_demo_date", lambda: date(2026, 8, 16))

    with pytest.raises(DemoBundleError, match="authority digest"):
        load_demo_bundle(
            root,
            expected_manifest_sha256=digest,
            expected_authority_sha256="0" * 64,
            expected_authority_classification="SYNTHETIC_TEST_ONLY",
            repo_root=REPO_ROOT,
        )

    authority_path = root / "authority.json"
    authority = json.loads(authority_path.read_text())
    authority["status"] = "PENDING"
    authority_path.write_text(
        json.dumps(authority, ensure_ascii=False, separators=(",", ":")) + "\n"
    )
    with pytest.raises(DemoBundleError, match="authority.status"):
        _load_bundle(root, digest)

    authority["status"] = "APPROVED"
    authority["decision_sha256"] = "0" * 64
    authority_path.write_text(
        json.dumps(authority, ensure_ascii=False, separators=(",", ":")) + "\n"
    )
    with pytest.raises(DemoBundleError, match="decision digest"):
        _load_bundle(root, digest)

    authority["decision_sha256"] = _sha256(
        (REPO_ROOT / authority["decision_ref"]).read_bytes()
    )
    authority["source_digests"][0]["sha256"] = "0" * 64
    authority_path.write_text(
        json.dumps(authority, ensure_ascii=False, separators=(",", ":")) + "\n"
    )
    with pytest.raises(DemoBundleError, match="source digests"):
        _load_bundle(root, digest)


def test_integrated_bundle_returns_exact_immutable_descriptors(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
):
    root, manifest, digest = _valid_integrated_bundle(tmp_path)
    monkeypatch.setattr(demo_bundle, "_current_demo_date", lambda: date(2026, 8, 21))

    snapshot = _load_bundle(root, digest)

    assert snapshot.schema_version == "fpms.demo-input-bundle/integrated-a-v1"
    assert snapshot.bundle_id == "fpms-integrated-a"
    assert snapshot.evidence_roles == tuple(INTEGRATED_EVIDENCE_ROLES)
    assert len(snapshot.evidence) == 12
    assert [row.role for row in snapshot.evidence] == INTEGRATED_EVIDENCE_ROLES
    assert [row.sha256 for row in snapshot.evidence] == [
        row["sha256"] for row in manifest["evidence"]
    ]
    assert snapshot.evidence[8].metadata.oa_sequence == 2
    assert snapshot.evidence[8].metadata.source_template_code == "DEMO_OA_NOTICE_2"
    assert (
        snapshot.evidence[11].metadata.supersedes_role
        == "GRANT_NOTICE_ORIGINAL"
    )
    with pytest.raises(AttributeError):
        snapshot.evidence[0].role = "MUTATED"

    authority = json.loads((root / "authority.json").read_text())
    assert len(authority["file_digests"]) == 13


def test_integrated_schema_is_additive_and_does_not_reinterpret_v1(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
):
    legacy_root, _legacy_manifest, legacy_digest = _valid_bundle(tmp_path / "legacy")
    integrated_root, _integrated_manifest, integrated_digest = _valid_integrated_bundle(
        tmp_path / "integrated"
    )
    monkeypatch.setattr(demo_bundle, "_current_demo_date", lambda: date(2026, 8, 21))

    legacy = _load_bundle(legacy_root, legacy_digest)
    integrated = _load_bundle(integrated_root, integrated_digest)

    assert legacy.schema_version == "fpms.demo-input-bundle/v1"
    assert legacy.evidence_roles == tuple(EVIDENCE_ROLES)
    assert integrated.schema_version == "fpms.demo-input-bundle/integrated-a-v1"
    assert integrated.evidence_roles == tuple(INTEGRATED_EVIDENCE_ROLES)


def test_integrated_roles_missing_extra_alias_and_schema_confusion_fail_closed(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
):
    monkeypatch.setattr(demo_bundle, "_current_demo_date", lambda: date(2026, 8, 21))

    root, manifest, _digest = _valid_integrated_bundle(tmp_path / "missing")
    manifest["evidence"].pop()
    with pytest.raises(DemoBundleError, match="roles or order"):
        _load_bundle(root, _write_manifest(root, manifest))

    root, manifest, _digest = _valid_integrated_bundle(tmp_path / "extra")
    manifest["evidence"].append(dict(manifest["evidence"][-1]))
    with pytest.raises(DemoBundleError, match="roles or order"):
        _load_bundle(root, _write_manifest(root, manifest))

    root, manifest, _digest = _valid_integrated_bundle(tmp_path / "alias")
    manifest["evidence"][6]["role"] = "OA_NOTICE"
    with pytest.raises(DemoBundleError, match="roles or order"):
        _load_bundle(root, _write_manifest(root, manifest))

    root, manifest, _digest = _valid_bundle(tmp_path / "schema-confusion")
    manifest["schema_version"] = "fpms.demo-input-bundle/integrated-a-v1"
    manifest["purpose"] = "LOCAL_INTEGRATED_A_E2E"
    manifest["contract"] = {
        "ref": INTEGRATED_CONTRACT_REF,
        "sha256": _sha256((REPO_ROOT / INTEGRATED_CONTRACT_REF).read_bytes()),
    }
    with pytest.raises(DemoBundleError, match="roles or order"):
        _load_bundle(root, _write_manifest(root, manifest))

    root, manifest, _digest = _valid_integrated_bundle(tmp_path / "unknown-schema")
    manifest["schema_version"] = "fpms.demo-input-bundle/integrated-a-v2"
    with pytest.raises(DemoBundleError, match="schema_version"):
        _load_bundle(root, _write_manifest(root, manifest))


def test_integrated_critical_hash_and_semantic_fallback_fail_closed(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
):
    monkeypatch.setattr(demo_bundle, "_current_demo_date", lambda: date(2026, 8, 21))

    root, manifest, _digest = _valid_integrated_bundle(tmp_path / "hash")
    first = manifest["evidence"][6]
    second = manifest["evidence"][8]
    first_bytes = (root / first["path"]).read_bytes()
    (root / second["path"]).write_bytes(first_bytes)
    second["size_bytes"] = len(first_bytes)
    second["sha256"] = first["sha256"]
    with pytest.raises(DemoBundleError, match="critical evidence hashes"):
        _load_bundle(root, _write_manifest(root, manifest))

    root, manifest, _digest = _valid_integrated_bundle(tmp_path / "oa-fallback")
    manifest["evidence"][8]["metadata"]["official_due_date"] = manifest["evidence"][6][
        "metadata"
    ]["official_due_date"]
    with pytest.raises(DemoBundleError, match="OA2 semantic metadata"):
        _load_bundle(root, _write_manifest(root, manifest))

    root, manifest, _digest = _valid_integrated_bundle(tmp_path / "grant-lineage")
    manifest["evidence"][11]["metadata"]["supersedes_role"] = "OA_NOTICE_2"
    with pytest.raises(DemoBundleError, match="supersedes_role"):
        _load_bundle(root, _write_manifest(root, manifest))


def test_integrated_authority_requires_exactly_thirteen_file_digests(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
):
    root, _manifest, digest = _valid_integrated_bundle(tmp_path)
    monkeypatch.setattr(demo_bundle, "_current_demo_date", lambda: date(2026, 8, 21))
    authority_path = root / "authority.json"
    authority = json.loads(authority_path.read_text())
    assert len(authority["file_digests"]) == 13
    authority["file_digests"].pop()
    authority_path.write_text(
        json.dumps(authority, ensure_ascii=False, separators=(",", ":")) + "\n"
    )
    with pytest.raises(DemoBundleError, match="file digests"):
        _load_bundle(root, digest)
