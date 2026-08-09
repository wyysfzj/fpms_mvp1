from __future__ import annotations

import hashlib
import importlib
from io import BytesIO
from types import MappingProxyType

import pytest
from docx import Document

from app.modules.documents.letter_context import (
    FormatLetterContextResult,
    FormatLetterNoticeVariant,
)


def _context_result(
    template_file_path: str = "templates/format_letters/format_letter_002.docx",
) -> FormatLetterContextResult:
    context = MappingProxyType(
        {
            "salutation_text": "尊敬的王&李<女士>：您好",
            "client_reference_no": "CLIENT-REF-001",
            "case_no": "23F0207CN",
            "invention_title": "一种 A&B <测试> 装置",
            "application_no": "CN202610000001",
            "filing_date_text": "2026-01-02",
            "applicant_names_text": "中国&辐射<防护>研究院",
            "inventor_names_text": "",
            "source_notice_name": "初步审查合格",
            "notice_variant_code": "PRELIMINARY_PASS",
            "publication_date_text": "",
            "deadline_text": "",
            "amount_lines_text": "",
            "template_variant_code": "FORMAT_LETTER_002",
        }
    )
    return FormatLetterContextResult(
        case_id="case-1",
        source_document_id="source-document-1",
        source_evidence_version_id="source-evidence-1",
        mapping_id="mapping-1",
        template_id="template-1",
        template_family_code="FORMAT_LETTER",
        template_variant_code="FORMAT_LETTER_002",
        template_file_path=template_file_path,
        notice_variant=FormatLetterNoticeVariant.PRELIMINARY_PASS,
        selected_contact_id="contact-1",
        contact_selection_source="EXPLICIT",
        salutation_source="SELECTED_CONTACT",
        context=context,
    )


def test_render_real_readable_word_with_required_name_and_content_hash() -> None:
    render_module = importlib.import_module("app.modules.documents.letter_render_service")

    rendered = render_module.render_format_letter(_context_result())

    assert rendered.file_name == "23F0207CN-给中国&辐射<防护>研究院的邮件.docx"
    assert (
        rendered.media_type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert rendered.content.startswith(b"PK")
    assert rendered.content_hash == f"sha256:{hashlib.sha256(rendered.content).hexdigest()}"

    document = Document(BytesIO(rendered.content))
    readable_text = "\n".join(
        [
            *(paragraph.text for paragraph in document.paragraphs),
            *(
                paragraph.text
                for table in document.tables
                for row in table.rows
                for cell in row.cells
                for paragraph in cell.paragraphs
            ),
        ]
    )
    assert "尊敬的王&李<女士>：您好" in readable_text
    assert "23F0207CN" in readable_text
    assert "一种 A&B <测试> 装置" in readable_text
    assert "CN202610000001" in readable_text
    assert "中国&辐射<防护>研究院" in readable_text
    assert "{{" not in readable_text


@pytest.mark.parametrize(
    "template_file_path",
    (
        "/templates/format_letters/format_letter_002.docx",
        "templates/format_letters/../format_letter_002.docx",
        "attachments",
        "templates/format_letters/missing.docx",
        "templates/format_letters",
    ),
)
def test_rejects_untrusted_or_non_file_template_paths(template_file_path: str) -> None:
    render_module = importlib.import_module("app.modules.documents.letter_render_service")

    with pytest.raises(ValueError, match="template_file_path"):
        render_module.render_format_letter(_context_result(template_file_path))
