#!/usr/bin/env python
"""Seed development database with default roles, permissions, and admin user."""

import json
import sys
import zipfile
from datetime import date
from decimal import Decimal
from pathlib import Path
from typing import NamedTuple
from uuid import uuid4

from docx import Document as DocxDocument
from docxtpl import DocxTemplate
from sqlalchemy.orm import Session

BASE_DIR = Path(__file__).resolve().parents[1]
if str(BASE_DIR) not in sys.path:
    sys.path.insert(0, str(BASE_DIR))

from app.core.security import get_password_hash  # noqa: E402
from app.db.session import SessionLocal  # noqa: E402
from app.models.system_param import SystemParam  # noqa: E402
from app.modules.auth.models import T_Role, T_User, T_UserRole  # noqa: E402
from app.modules.cases.models import Case, T_CaseApplicant, T_CaseInventor  # noqa: E402
from app.modules.documents.models import DocTemplate  # noqa: E402
from app.modules.documents.official_notice_catalog import (  # noqa: E402
    seed_fee_reduction_approval_official_notice_catalog,
)
from app.modules.documents.official_notice_catalog import (  # noqa: E402
    seed_official_letter_out_form_001_catalog as seed_official_letter_out_catalog,
)
from app.modules.fees.models import FeeRate  # noqa: E402
from app.modules.masterdata.applicants.models import Applicant  # noqa: E402
from app.modules.masterdata.clients.models import Client  # noqa: E402
from app.modules.rbac.service import seed_default_roles_perms  # noqa: E402
from app.modules.tasks.models import TaskTemplate  # noqa: E402
from app.modules.templates.models import Template  # noqa: E402

GRANT_FEE_NOTICE_TEMPLATE_CODE = "GRANT_FEE_NOTICE"
GRANT_FEE_NOTICE_TEMPLATE_PATH = "templates/grant_fee_notice.docx"
OFFICIAL_FEE_SOURCE_DOC = "docs/postdemo/专利收费场景-20260626.docx"
OFFICIAL_FEE_SOURCE_URL = "http://www.tianyueip.com/product/612"
OFFICIAL_FEE_SOURCE_POLICY = "客户收费场景DOCX + 天悦网页线索；上线执行前需客户或官方来源确认"
OFFICIAL_FEE_SOURCE_VERSION = "2026-07-05-postdemo"


def _json_params(value: dict) -> str:
    return json.dumps(value, ensure_ascii=False, sort_keys=True)


def _official_fee_rate_catalog() -> list[dict]:
    """Return the post-demo official-fee parameter catalog."""
    confirmed = "CONFIRMED"
    pending = "PENDING_CONFIRMATION"
    domestic_source = {
        "fee_type": "GOV",
        "currency": "CNY",
        "country_code": "CN",
        "fee_domain": "PATENT",
        "fee_section": "专利收费-国内部分",
        "source_doc": OFFICIAL_FEE_SOURCE_DOC,
        "source_url": OFFICIAL_FEE_SOURCE_URL,
        "source_policy": OFFICIAL_FEE_SOURCE_POLICY,
        "source_version": OFFICIAL_FEE_SOURCE_VERSION,
    }
    annuity_reduction_note = {
        "allow_reduction_years_from_grant": 10,
        "late_fee_reduction": False,
    }

    return [
        {
            **domestic_source,
            "fee_code": "CN_INV_APPLICATION_FEE",
            "fee_name": "发明专利申请费",
            "default_amount": Decimal("900.00"),
            "enabled": True,
            "rate_group": "DOMESTIC",
            "patent_category": "INV",
            "fee_category": "申请费",
            "fee_subtype": "发明专利",
            "reduction_scope": "申请费（不包括公布印刷费、申请附加费）",
            "calc_mode": "FIXED",
            "allow_reduction": True,
            "source_status": confirmed,
        },
        {
            **domestic_source,
            "fee_code": "CN_UM_APPLICATION_FEE",
            "fee_name": "实用新型专利申请费",
            "default_amount": Decimal("500.00"),
            "enabled": True,
            "rate_group": "DOMESTIC",
            "patent_category": "UM",
            "fee_category": "申请费",
            "fee_subtype": "实用新型专利",
            "reduction_scope": "申请费（不包括公布印刷费、申请附加费）",
            "calc_mode": "FIXED",
            "allow_reduction": True,
            "source_status": confirmed,
        },
        {
            **domestic_source,
            "fee_code": "CN_DES_APPLICATION_FEE",
            "fee_name": "外观设计专利申请费",
            "default_amount": Decimal("500.00"),
            "enabled": True,
            "rate_group": "DOMESTIC",
            "patent_category": "DES",
            "fee_category": "申请费",
            "fee_subtype": "外观设计专利",
            "reduction_scope": "申请费（不包括公布印刷费、申请附加费）",
            "calc_mode": "FIXED",
            "allow_reduction": True,
            "source_status": confirmed,
        },
        {
            **domestic_source,
            "fee_code": "CN_EXCESS_CLAIM_FEE",
            "fee_name": "权利要求附加费（第11项起每项）",
            "default_amount": Decimal("150.00"),
            "enabled": True,
            "rate_group": "DOMESTIC",
            "calc_mode": "PER_CLAIM",
            "calc_params": _json_params({"threshold": 10, "unit_amount": "150.00"}),
            "allow_reduction": False,
            "source_status": confirmed,
        },
        {
            **domestic_source,
            "fee_code": "CN_SPEC_PAGE_31_300_FEE",
            "fee_name": "说明书附加费（31-300页每页）",
            "default_amount": Decimal("50.00"),
            "enabled": False,
            "rate_group": "DOMESTIC",
            "calc_mode": "PER_PAGE",
            "calc_params": _json_params({"from_page": 31, "to_page": 300, "unit_amount": "50.00"}),
            "allow_reduction": False,
            "source_status": pending,
        },
        {
            **domestic_source,
            "fee_code": "CN_SPEC_PAGE_301_PLUS_FEE",
            "fee_name": "说明书附加费（301页起每页）",
            "default_amount": Decimal("100.00"),
            "enabled": False,
            "rate_group": "DOMESTIC",
            "calc_mode": "PER_PAGE",
            "calc_params": _json_params({"from_page": 301, "unit_amount": "100.00"}),
            "allow_reduction": False,
            "source_status": pending,
        },
        {
            **domestic_source,
            "fee_code": "CN_PUBLICATION_PRINT_FEE",
            "fee_name": "公布印刷费",
            "default_amount": Decimal("50.00"),
            "enabled": True,
            "rate_group": "DOMESTIC",
            "patent_category": "INV",
            "fee_category": "公布印刷费",
            "fee_subtype": "仅发明专利",
            "reduction_scope": "不可费减",
            "calc_mode": "FIXED",
            "allow_reduction": False,
            "source_status": confirmed,
        },
        {
            **domestic_source,
            "fee_code": "CN_PRIORITY_CLAIM_FEE",
            "fee_name": "优先权要求费（每项）",
            "default_amount": Decimal("80.00"),
            "enabled": False,
            "rate_group": "DOMESTIC",
            "calc_mode": "COMPOSITE",
            "calc_params": _json_params({"unit": "priority", "unit_amount": "80.00"}),
            "allow_reduction": False,
            "source_status": pending,
        },
        {
            **domestic_source,
            "fee_code": "CN_SUBSTANTIVE_EXAM_FEE",
            "fee_name": "发明专利实质审查费",
            "default_amount": Decimal("2500.00"),
            "enabled": True,
            "rate_group": "DOMESTIC",
            "patent_category": "INV",
            "calc_mode": "FIXED",
            "allow_reduction": True,
            "source_status": confirmed,
        },
        {
            **domestic_source,
            "fee_code": "CN_REEXAM_FEE_INV",
            "fee_name": "复审费（发明）",
            "default_amount": Decimal("1000.00"),
            "enabled": False,
            "rate_group": "DOMESTIC",
            "patent_category": "INV",
            "calc_mode": "FIXED",
            "allow_reduction": True,
            "source_status": pending,
        },
        {
            **domestic_source,
            "fee_code": "CN_REEXAM_FEE_UM",
            "fee_name": "复审费（实用新型）",
            "default_amount": Decimal("300.00"),
            "enabled": False,
            "rate_group": "DOMESTIC",
            "patent_category": "UM",
            "fee_category": "复审费",
            "fee_subtype": "实用新型专利",
            "reduction_scope": "复审费",
            "calc_mode": "FIXED",
            "allow_reduction": True,
            "source_status": pending,
        },
        {
            **domestic_source,
            "fee_code": "CN_REEXAM_FEE_DES",
            "fee_name": "复审费（外观设计）",
            "default_amount": Decimal("300.00"),
            "enabled": False,
            "rate_group": "DOMESTIC",
            "patent_category": "DES",
            "fee_category": "复审费",
            "fee_subtype": "外观设计专利",
            "reduction_scope": "复审费",
            "calc_mode": "FIXED",
            "allow_reduction": True,
            "source_status": pending,
        },
        {
            **domestic_source,
            "fee_code": "CN_ANNUITY_FEE_INV",
            "fee_name": "发明专利年费",
            "default_amount": Decimal("0.00"),
            "enabled": True,
            "rate_group": "ANNUITY",
            "patent_category": "INV",
            "calc_mode": "TIER",
            "calc_params": _json_params(
                {
                    **annuity_reduction_note,
                    "tiers": [
                        {"from": 1, "to": 3, "amount": "900.00"},
                        {"from": 4, "to": 6, "amount": "1200.00"},
                        {"from": 7, "to": 9, "amount": "2000.00"},
                        {"from": 10, "to": 12, "amount": "4000.00"},
                        {"from": 13, "to": 15, "amount": "6000.00"},
                        {"from": 16, "to": 20, "amount": "8000.00"},
                    ],
                }
            ),
            "allow_reduction": True,
            "source_status": confirmed,
        },
        {
            **domestic_source,
            "fee_code": "CN_ANNUITY_FEE_UM",
            "fee_name": "实用新型专利年费",
            "default_amount": Decimal("0.00"),
            "enabled": True,
            "rate_group": "ANNUITY",
            "patent_category": "UM",
            "calc_mode": "TIER",
            "calc_params": _json_params(
                {
                    **annuity_reduction_note,
                    "tiers": [
                        {"from": 1, "to": 3, "amount": "600.00"},
                        {"from": 4, "to": 5, "amount": "900.00"},
                        {"from": 6, "to": 8, "amount": "1200.00"},
                        {"from": 9, "to": 10, "amount": "2000.00"},
                    ],
                }
            ),
            "allow_reduction": True,
            "source_status": confirmed,
        },
        {
            **domestic_source,
            "fee_code": "CN_ANNUITY_FEE_DES",
            "fee_name": "外观设计专利年费",
            "default_amount": Decimal("0.00"),
            "enabled": True,
            "rate_group": "ANNUITY",
            "patent_category": "DES",
            "calc_mode": "TIER",
            "calc_params": _json_params(
                {
                    **annuity_reduction_note,
                    "tiers": [
                        {"from": 1, "to": 3, "amount": "600.00"},
                        {"from": 4, "to": 5, "amount": "900.00"},
                        {"from": 6, "to": 8, "amount": "1200.00"},
                        {"from": 9, "to": 10, "amount": "2000.00"},
                        {"from": 11, "to": 15, "amount": "3000.00"},
                    ],
                }
            ),
            "allow_reduction": True,
            "source_status": confirmed,
        },
        {
            **domestic_source,
            "fee_code": "CN_ANNUITY_LATE_FEE",
            "fee_name": "年费滞纳金",
            "default_amount": Decimal("0.00"),
            "enabled": False,
            "rate_group": "ANNUITY",
            "calc_mode": "COMPOSITE",
            "calc_params": _json_params({"monthly_percent": "5", "max_percent": "25"}),
            "allow_reduction": False,
            "source_status": pending,
        },
        {
            **domestic_source,
            "fee_code": "CN_RESTORE_RIGHT_FEE",
            "fee_name": "恢复权利请求费",
            "default_amount": Decimal("1000.00"),
            "enabled": False,
            "rate_group": "DOMESTIC",
            "calc_mode": "FIXED",
            "allow_reduction": False,
            "source_status": pending,
        },
        {
            **domestic_source,
            "fee_code": "CN_EXTENSION_FIRST_MONTH_FEE",
            "fee_name": "第一次延长期限请求费（每月）",
            "default_amount": Decimal("300.00"),
            "enabled": False,
            "rate_group": "DOMESTIC",
            "calc_mode": "FIXED",
            "calc_params": _json_params({"unit": "month"}),
            "allow_reduction": False,
            "source_status": pending,
        },
        {
            **domestic_source,
            "fee_code": "CN_EXTENSION_REPEAT_MONTH_FEE",
            "fee_name": "再次延长期限请求费（每月）",
            "default_amount": Decimal("2000.00"),
            "enabled": False,
            "rate_group": "DOMESTIC",
            "calc_mode": "FIXED",
            "calc_params": _json_params({"unit": "month"}),
            "allow_reduction": False,
            "source_status": pending,
        },
        {
            **domestic_source,
            "fee_code": "CN_BIBLIO_CHANGE_FEE",
            "fee_name": "著录事项变更费",
            "default_amount": Decimal("200.00"),
            "enabled": False,
            "rate_group": "DOMESTIC",
            "calc_mode": "FIXED",
            "allow_reduction": False,
            "source_status": pending,
        },
        {
            **domestic_source,
            "fee_code": "CN_PATENT_EVALUATION_REPORT_UM",
            "fee_name": "专利权评价报告请求费（实用新型）",
            "default_amount": Decimal("2400.00"),
            "enabled": False,
            "rate_group": "DOMESTIC",
            "patent_category": "UM",
            "calc_mode": "FIXED",
            "allow_reduction": False,
            "source_status": pending,
        },
        {
            **domestic_source,
            "fee_code": "CN_PATENT_EVALUATION_REPORT_DES",
            "fee_name": "专利权评价报告请求费（外观设计）",
            "default_amount": Decimal("2400.00"),
            "enabled": False,
            "rate_group": "DOMESTIC",
            "patent_category": "DES",
            "calc_mode": "FIXED",
            "allow_reduction": False,
            "source_status": pending,
        },
        {
            **domestic_source,
            "fee_code": "CN_INVALIDATION_REQUEST_INV",
            "fee_name": "无效宣告请求费（发明）",
            "default_amount": Decimal("3000.00"),
            "enabled": False,
            "rate_group": "DOMESTIC",
            "patent_category": "INV",
            "calc_mode": "FIXED",
            "allow_reduction": False,
            "source_status": pending,
        },
        {
            **domestic_source,
            "fee_code": "CN_INVALIDATION_REQUEST_UM",
            "fee_name": "无效宣告请求费（实用新型）",
            "default_amount": Decimal("1500.00"),
            "enabled": False,
            "rate_group": "DOMESTIC",
            "patent_category": "UM",
            "calc_mode": "FIXED",
            "allow_reduction": False,
            "source_status": pending,
        },
        {
            **domestic_source,
            "fee_code": "CN_INVALIDATION_REQUEST_DES",
            "fee_name": "无效宣告请求费（外观设计）",
            "default_amount": Decimal("1500.00"),
            "enabled": False,
            "rate_group": "DOMESTIC",
            "patent_category": "DES",
            "calc_mode": "FIXED",
            "allow_reduction": False,
            "source_status": pending,
        },
        {
            **domestic_source,
            "fee_code": "CN_FILE_COPY_CERT_FEE",
            "fee_name": "专利文件副本证明费（每份）",
            "default_amount": Decimal("30.00"),
            "enabled": False,
            "rate_group": "DOMESTIC",
            "calc_mode": "FIXED",
            "allow_reduction": False,
            "source_status": pending,
        },
        {
            **domestic_source,
            "fee_code": "PCT_SEARCH_FEE",
            "fee_name": "PCT国际阶段检索费",
            "default_amount": Decimal("2100.00"),
            "enabled": False,
            "rate_group": "PCT",
            "case_type": "PCT_INTL",
            "fee_section": "专利收费-PCT申请收费",
            "fee_category": "PCT 国际阶段费用",
            "fee_subtype": "检索费",
            "reduction_scope": "不可费减",
            "calc_mode": "FIXED",
            "allow_reduction": False,
            "source_status": pending,
        },
        {
            **domestic_source,
            "fee_code": "PCT_ADDITIONAL_SEARCH_FEE",
            "fee_name": "PCT附加检索费（每个单一性主题）",
            "default_amount": Decimal("2100.00"),
            "enabled": False,
            "rate_group": "PCT",
            "case_type": "PCT_INTL",
            "calc_mode": "COMPOSITE",
            "calc_params": _json_params({"unit": "unity_subject"}),
            "allow_reduction": False,
            "source_status": pending,
        },
        {
            **domestic_source,
            "fee_code": "PCT_PRIORITY_DOC_FEE",
            "fee_name": "PCT优先权文件费（每项）",
            "default_amount": Decimal("150.00"),
            "enabled": False,
            "rate_group": "PCT",
            "case_type": "PCT_INTL",
            "calc_mode": "COMPOSITE",
            "calc_params": _json_params({"unit": "priority"}),
            "allow_reduction": False,
            "source_status": pending,
        },
        {
            **domestic_source,
            "fee_code": "PCT_PRELIM_EXAM_FEE",
            "fee_name": "PCT初步审查费",
            "default_amount": Decimal("1500.00"),
            "enabled": False,
            "rate_group": "PCT",
            "case_type": "PCT_INTL",
            "calc_mode": "FIXED",
            "allow_reduction": False,
            "source_status": pending,
        },
        {
            **domestic_source,
            "fee_code": "PCT_PRELIM_EXAM_ADDITIONAL_FEE",
            "fee_name": "PCT初步审查附加费（每个主题）",
            "default_amount": Decimal("1500.00"),
            "enabled": False,
            "rate_group": "PCT",
            "case_type": "PCT_INTL",
            "calc_mode": "COMPOSITE",
            "calc_params": _json_params({"unit": "unity_subject"}),
            "allow_reduction": False,
            "source_status": pending,
        },
        {
            **domestic_source,
            "fee_code": "PCT_UNITY_OBJECTION_FEE",
            "fee_name": "PCT单一性异议费",
            "default_amount": Decimal("200.00"),
            "enabled": False,
            "rate_group": "PCT",
            "case_type": "PCT_INTL",
            "calc_mode": "FIXED",
            "allow_reduction": False,
            "source_status": pending,
        },
        {
            **domestic_source,
            "fee_code": "PCT_COPY_FEE_PER_PAGE",
            "fee_name": "PCT副本复制费（每页）",
            "default_amount": Decimal("2.00"),
            "enabled": False,
            "rate_group": "PCT",
            "case_type": "PCT_INTL",
            "calc_mode": "PER_PAGE",
            "allow_reduction": False,
            "source_status": pending,
        },
        {
            **domestic_source,
            "fee_code": "PCT_LATE_FILING_FEE",
            "fee_name": "PCT后提交费",
            "default_amount": Decimal("200.00"),
            "enabled": False,
            "rate_group": "PCT",
            "case_type": "PCT_INTL",
            "calc_mode": "FIXED",
            "allow_reduction": False,
            "source_status": pending,
        },
        {
            **domestic_source,
            "fee_code": "PCT_RESTORE_RIGHT_FEE",
            "fee_name": "PCT国际阶段恢复权利请求费",
            "default_amount": Decimal("1000.00"),
            "enabled": False,
            "rate_group": "PCT",
            "case_type": "PCT_INTL",
            "calc_mode": "FIXED",
            "allow_reduction": False,
            "source_status": pending,
        },
        {
            **domestic_source,
            "fee_code": "PCT_INTL_LATE_FEE",
            "fee_name": "PCT国际阶段滞纳金",
            "default_amount": Decimal("0.00"),
            "enabled": False,
            "rate_group": "PCT",
            "case_type": "PCT_INTL",
            "calc_mode": "COMPOSITE",
            "calc_params": _json_params({"unpaid_fee_percent": "50", "cap_rule": "国际申请费50%"}),
            "allow_reduction": False,
            "source_status": pending,
        },
        {
            **domestic_source,
            "fee_code": "PCT_CN_GRACE_FEE",
            "fee_name": "PCT进入中国国家阶段宽限费",
            "default_amount": Decimal("1000.00"),
            "enabled": False,
            "rate_group": "PCT_CN",
            "case_type": "PCT_NATL",
            "calc_mode": "FIXED",
            "allow_reduction": False,
            "source_status": pending,
        },
        {
            **domestic_source,
            "fee_code": "PCT_CN_TRANSLATION_CORRECTION_PRELIM_FEE",
            "fee_name": "PCT进入中国译文改正费（初审阶段）",
            "default_amount": Decimal("300.00"),
            "enabled": False,
            "rate_group": "PCT_CN",
            "case_type": "PCT_NATL",
            "calc_mode": "FIXED",
            "allow_reduction": False,
            "source_status": pending,
        },
        {
            **domestic_source,
            "fee_code": "PCT_CN_TRANSLATION_CORRECTION_EXAM_FEE",
            "fee_name": "PCT进入中国译文改正费（实审阶段）",
            "default_amount": Decimal("1200.00"),
            "enabled": False,
            "rate_group": "PCT_CN",
            "case_type": "PCT_NATL",
            "calc_mode": "FIXED",
            "allow_reduction": False,
            "source_status": pending,
        },
        {
            **domestic_source,
            "fee_code": "PCT_CN_UNITY_RESTORE_FEE",
            "fee_name": "PCT进入中国单一性恢复费",
            "default_amount": Decimal("900.00"),
            "enabled": False,
            "rate_group": "PCT_CN",
            "case_type": "PCT_NATL",
            "calc_mode": "FIXED",
            "allow_reduction": False,
            "source_status": pending,
        },
        {
            **domestic_source,
            "fee_code": "PCT_CN_PRIORITY_RESTORE_FEE",
            "fee_name": "PCT进入中国优先权恢复费",
            "default_amount": Decimal("1000.00"),
            "enabled": False,
            "rate_group": "PCT_CN",
            "case_type": "PCT_NATL",
            "calc_mode": "FIXED",
            "allow_reduction": False,
            "source_status": pending,
        },
        {
            **domestic_source,
            "fee_code": "CN_PATENT_TERM_COMPENSATION_REQUEST_FEE",
            "fee_name": "专利权期限补偿请求费",
            "default_amount": Decimal("200.00"),
            "enabled": False,
            "rate_group": "COMPENSATION",
            "fee_category": "专利权期限补偿请求费",
            "fee_subtype": "每件",
            "reduction_scope": "不可费减",
            "calc_mode": "FIXED",
            "allow_reduction": False,
            "source_status": pending,
        },
        {
            **domestic_source,
            "fee_code": "CN_COMPENSATION_PERIOD_ANNUITY_FEE",
            "fee_name": "专利权补偿期年费",
            "default_amount": Decimal("8000.00"),
            "enabled": False,
            "rate_group": "COMPENSATION",
            "fee_category": "专利权补偿期年费",
            "fee_subtype": "每年，不足一年部分不收取",
            "reduction_scope": "不可费减",
            "calc_mode": "BY_YEAR",
            "allow_reduction": False,
            "source_status": pending,
        },
        {
            **domestic_source,
            "fee_code": "CN_HAGUE_DESIGN_CN_DESIGNATION_FEE_FIRST",
            "fee_name": "外观设计国际注册指定中国单独指定费（第一期）",
            "default_amount": Decimal("4100.00"),
            "enabled": False,
            "rate_group": "HAGUE",
            "patent_category": "DES",
            "fee_section": "专利收费-外观设计国际注册申请",
            "fee_category": "指定中国单独指定费",
            "fee_subtype": "第一期",
            "reduction_scope": "指定中国单独指定费第一期",
            "calc_mode": "FIXED",
            "allow_reduction": True,
            "source_status": pending,
        },
        {
            **domestic_source,
            "fee_code": "CN_HAGUE_DESIGN_CN_DESIGNATION_FEE_SECOND",
            "fee_name": "外观设计国际注册指定中国单独指定费（第二期）",
            "default_amount": Decimal("7600.00"),
            "enabled": False,
            "rate_group": "HAGUE",
            "patent_category": "DES",
            "fee_section": "专利收费-外观设计国际注册申请",
            "fee_category": "指定中国单独指定费",
            "fee_subtype": "第二期",
            "reduction_scope": "指定中国单独指定费第二期",
            "calc_mode": "FIXED",
            "allow_reduction": True,
            "source_status": pending,
        },
        {
            **domestic_source,
            "fee_code": "CN_HAGUE_DESIGN_CN_DESIGNATION_FEE",
            "fee_name": "外观设计国际注册指定中国单独指定费",
            "default_amount": Decimal("15000.00"),
            "enabled": False,
            "rate_group": "HAGUE",
            "patent_category": "DES",
            "fee_section": "专利收费-外观设计国际注册申请",
            "fee_category": "指定中国单独指定费",
            "fee_subtype": "第三期",
            "reduction_scope": "不可费减",
            "calc_mode": "FIXED",
            "allow_reduction": False,
            "source_status": pending,
        },
        {
            **domestic_source,
            "fee_code": "IC_LAYOUT_REGISTRATION_FEE",
            "fee_name": "布图设计登记费（每件）",
            "default_amount": Decimal("1000.00"),
            "enabled": False,
            "rate_group": "IC_LAYOUT",
            "fee_domain": "IC_LAYOUT",
            "fee_section": "集成电路布图设计收费标准",
            "fee_category": "布图设计登记费",
            "fee_subtype": "每件",
            "reduction_scope": "不可费减",
            "calc_mode": "FIXED",
            "allow_reduction": False,
            "source_status": pending,
        },
        {
            **domestic_source,
            "fee_code": "IC_LAYOUT_REEXAM_REQUEST_FEE",
            "fee_name": "布图设计登记复审请求费（每件）",
            "default_amount": Decimal("1000.00"),
            "enabled": False,
            "rate_group": "IC_LAYOUT",
            "fee_domain": "IC_LAYOUT",
            "fee_section": "集成电路布图设计收费标准",
            "fee_category": "布图设计登记复审请求费",
            "fee_subtype": "每件",
            "reduction_scope": "不可费减",
            "calc_mode": "FIXED",
            "allow_reduction": False,
            "source_status": pending,
        },
        {
            **domestic_source,
            "fee_code": "IC_LAYOUT_BIBLIO_CHANGE_FEE",
            "fee_name": "布图设计著录事项变更手续费（每件每次）",
            "default_amount": Decimal("50.00"),
            "enabled": False,
            "rate_group": "IC_LAYOUT",
            "fee_domain": "IC_LAYOUT",
            "fee_section": "集成电路布图设计收费标准",
            "fee_category": "著录事项变更手续费",
            "fee_subtype": "每件每次",
            "reduction_scope": "不可费减",
            "calc_mode": "FIXED",
            "allow_reduction": False,
            "source_status": pending,
        },
        {
            **domestic_source,
            "fee_code": "IC_LAYOUT_EXTENSION_REQUEST_FEE",
            "fee_name": "布图设计延长期限请求费（每件每次）",
            "default_amount": Decimal("150.00"),
            "enabled": False,
            "rate_group": "IC_LAYOUT",
            "fee_domain": "IC_LAYOUT",
            "fee_section": "集成电路布图设计收费标准",
            "fee_category": "延长期限请求费",
            "fee_subtype": "每件每次",
            "reduction_scope": "不可费减",
            "calc_mode": "FIXED",
            "allow_reduction": False,
            "source_status": pending,
        },
        {
            **domestic_source,
            "fee_code": "IC_LAYOUT_RESTORE_RIGHT_FEE",
            "fee_name": "恢复布图设计登记权利请求费（每件）",
            "default_amount": Decimal("500.00"),
            "enabled": False,
            "rate_group": "IC_LAYOUT",
            "fee_domain": "IC_LAYOUT",
            "fee_section": "集成电路布图设计收费标准",
            "fee_category": "恢复布图设计登记权利请求费",
            "fee_subtype": "每件",
            "reduction_scope": "不可费减",
            "calc_mode": "FIXED",
            "allow_reduction": False,
            "source_status": pending,
        },
        {
            **domestic_source,
            "fee_code": "IC_LAYOUT_NONVOLUNTARY_LICENSE_REQUEST_FEE",
            "fee_name": "非自愿许可使用布图设计请求费（每件）",
            "default_amount": Decimal("150.00"),
            "enabled": False,
            "rate_group": "IC_LAYOUT",
            "fee_domain": "IC_LAYOUT",
            "fee_section": "集成电路布图设计收费标准",
            "fee_category": "非自愿许可使用布图设计请求费",
            "fee_subtype": "每件",
            "reduction_scope": "不可费减",
            "calc_mode": "FIXED",
            "allow_reduction": False,
            "source_status": pending,
        },
        {
            **domestic_source,
            "fee_code": "IC_LAYOUT_NONVOLUNTARY_LICENSE_REMUNERATION_ADJUDICATION_FEE",
            "fee_name": "非自愿许可使用布图设计支付报酬裁决费（每件）",
            "default_amount": Decimal("150.00"),
            "enabled": False,
            "rate_group": "IC_LAYOUT",
            "fee_domain": "IC_LAYOUT",
            "fee_section": "集成电路布图设计收费标准",
            "fee_category": "非自愿许可使用布图设计支付报酬裁决费",
            "fee_subtype": "每件",
            "reduction_scope": "不可费减",
            "calc_mode": "FIXED",
            "allow_reduction": False,
            "source_status": pending,
        },
    ]


def seed_admin_user(db: Session) -> None:
    """Create default admin user. Idempotent."""
    username = "admin"

    admin_role = db.query(T_Role).filter(T_Role.code == "Admin").first()
    if not admin_role:
        raise RuntimeError("Admin role 'Admin' not found. Seed roles first.")

    admin_user = db.query(T_User).filter(T_User.username == username).first()
    created_user = False
    if not admin_user:
        admin_user = T_User(
            id=str(uuid4()),
            username=username,
            display_name="Administrator",
            password_hash=get_password_hash("admin123"),
            is_active=True,
        )
        db.add(admin_user)
        db.flush()
        created_user = True

    existing_binding = (
        db.query(T_UserRole)
        .filter(
            T_UserRole.user_id == admin_user.id,
            T_UserRole.role_id == admin_role.id,
        )
        .first()
    )
    created_binding = False
    if not existing_binding:
        db.add(T_UserRole(user_id=admin_user.id, role_id=admin_role.id))
        created_binding = True

    db.commit()

    if created_user:
        print(f"Created admin user '{username}' with password 'admin123'")
    elif created_binding:
        print(f"Bound admin user '{username}' to Admin role")
    else:
        print(f"Admin user '{username}' already exists")


def seed_v3_cases(db: Session) -> None:
    """Create V3 workflow stepper test cases covering all 13 statuses. Idempotent."""
    # Check if V3 cases already seeded
    existing = db.query(Case).filter(Case.case_no == "V3-001").first()
    if existing:
        print("V3 test cases already exist, skipping")
        return

    # Create test clients
    clients_data = [
        {"code": "C-NIO", "name_cn": "蔚来汽车科技有限公司", "name_en": "NIO Inc."},
        {"code": "C-BYD", "name_cn": "比亚迪股份有限公司", "name_en": "BYD Company Limited"},
        {"code": "C-HW", "name_cn": "华为技术有限公司", "name_en": "Huawei Technologies Co., Ltd."},
        {"code": "C-XM", "name_cn": "小米科技有限责任公司", "name_en": "Xiaomi Corporation"},
    ]

    client_ids = {}
    for cd in clients_data:
        existing_client = db.query(Client).filter(Client.client_code == cd["code"]).first()
        if existing_client:
            client_ids[cd["code"]] = existing_client.id
        else:
            cid = str(uuid4())
            db.add(
                Client(
                    id=cid,
                    client_code=cd["code"],
                    name_cn=cd["name_cn"],
                    name_en=cd["name_en"],
                )
            )
            client_ids[cd["code"]] = cid
    db.flush()

    # V3 test cases — one per status, all in Simplified Chinese
    v3_cases = [
        {
            "case_no": "V3-001",
            "title_cn": "智能充电桩控制方法及系统",
            "title_en": "Smart Charging Pile Control Method and System",
            "app_no": "202410001001.1",
            "status": "WAITING_RECEIPT",
            "client_code": "C-NIO",
            "filing_date": date(2024, 1, 15),
            "recv_date": date(2024, 1, 10),
            "applicant": "蔚来汽车科技有限公司",
            "inventor": "张伟",
        },
        {
            "case_no": "V3-002",
            "title_cn": "电池热管理温控装置",
            "title_en": "Battery Thermal Management Temperature Control Device",
            "app_no": "202410001002.6",
            "status": "PRELIM_EXAM",
            "client_code": "C-NIO",
            "filing_date": date(2024, 2, 20),
            "recv_date": date(2024, 2, 15),
            "applicant": "蔚来汽车科技有限公司",
            "inventor": "李强",
        },
        {
            "case_no": "V3-003",
            "title_cn": "自动驾驶路径规划算法",
            "title_en": "Autonomous Driving Path Planning Algorithm",
            "app_no": "202410001003.0",
            "status": "PRELIM_PASS",
            "client_code": "C-BYD",
            "filing_date": date(2024, 3, 10),
            "recv_date": date(2024, 3, 5),
            "applicant": "比亚迪股份有限公司",
            "inventor": "王芳",
        },
        {
            "case_no": "V3-004",
            "title_cn": "车载激光雷达信号处理方法",
            "title_en": "Vehicle Lidar Signal Processing Method",
            "app_no": "202410001004.5",
            "status": "AMENDMENT",
            "client_code": "C-BYD",
            "filing_date": date(2024, 3, 25),
            "recv_date": date(2024, 3, 20),
            "applicant": "比亚迪股份有限公司",
            "inventor": "赵敏",
        },
        {
            "case_no": "V3-005",
            "title_cn": "5G基站天线阵列优化设计",
            "title_en": "5G Base Station Antenna Array Optimization Design",
            "app_no": "202410001005.X",
            "status": "PUBLISHED",
            "client_code": "C-HW",
            "filing_date": date(2024, 4, 12),
            "recv_date": date(2024, 4, 8),
            "applicant": "华为技术有限公司",
            "inventor": "陈刚",
        },
        {
            "case_no": "V3-006",
            "title_cn": "分布式数据库一致性协议",
            "title_en": "Distributed Database Consistency Protocol",
            "app_no": "202410001006.4",
            "status": "SUB_EXAM",
            "client_code": "C-HW",
            "filing_date": date(2024, 5, 8),
            "recv_date": date(2024, 5, 3),
            "applicant": "华为技术有限公司",
            "inventor": "刘洋",
        },
        {
            "case_no": "V3-007",
            "title_cn": "手机摄像模组光学防抖方法",
            "title_en": "Smartphone Camera Module OIS Method",
            "app_no": "202410001007.9",
            "status": "OA1",
            "client_code": "C-XM",
            "filing_date": date(2024, 5, 22),
            "recv_date": date(2024, 5, 18),
            "applicant": "小米科技有限责任公司",
            "inventor": "周磊",
        },
        {
            "case_no": "V3-008",
            "title_cn": "智能家居语音控制交互系统",
            "title_en": "Smart Home Voice Control Interaction System",
            "app_no": "202410001008.3",
            "status": "OA2",
            "client_code": "C-XM",
            "filing_date": date(2024, 6, 5),
            "recv_date": date(2024, 6, 1),
            "applicant": "小米科技有限责任公司",
            "inventor": "吴娜",
        },
        {
            "case_no": "V3-009",
            "title_cn": "芯片制造工艺缺陷检测方法",
            "title_en": "Chip Manufacturing Process Defect Detection Method",
            "app_no": "202410001009.8",
            "status": "REEXAM",
            "client_code": "C-HW",
            "filing_date": date(2024, 6, 20),
            "recv_date": date(2024, 6, 15),
            "applicant": "华为技术有限公司",
            "inventor": "孙涛",
        },
        {
            "case_no": "V3-010",
            "title_cn": "新能源汽车能量回收控制策略",
            "title_en": "New Energy Vehicle Energy Recovery Control Strategy",
            "app_no": "202410001010.0",
            "status": "GRANTED",
            "client_code": "C-BYD",
            "filing_date": date(2024, 1, 8),
            "recv_date": date(2024, 1, 3),
            "applicant": "比亚迪股份有限公司",
            "inventor": "郑华",
        },
        {
            "case_no": "V3-011",
            "title_cn": "无线充电效率提升装置",
            "title_en": "Wireless Charging Efficiency Enhancement Device",
            "app_no": "202410001011.5",
            "status": "REJECTED",
            "client_code": "C-XM",
            "filing_date": date(2024, 2, 12),
            "recv_date": date(2024, 2, 8),
            "applicant": "小米科技有限责任公司",
            "inventor": "马超",
        },
        {
            "case_no": "V3-012",
            "title_cn": "固态电池电解质制备方法",
            "title_en": "Solid-State Battery Electrolyte Preparation Method",
            "app_no": "202410001012.X",
            "status": "TERMINATED",
            "client_code": "C-NIO",
            "filing_date": date(2023, 6, 15),
            "recv_date": date(2023, 6, 10),
            "applicant": "蔚来汽车科技有限公司",
            "inventor": "黄丽",
        },
        {
            "case_no": "V3-013",
            "title_cn": "物联网设备安全认证协议",
            "title_en": "IoT Device Security Authentication Protocol",
            "app_no": "202410001013.4",
            "status": "INVALIDATED",
            "client_code": "C-HW",
            "filing_date": date(2023, 3, 20),
            "recv_date": date(2023, 3, 15),
            "applicant": "华为技术有限公司",
            "inventor": "林峰",
        },
    ]

    for c in v3_cases:
        case_id = str(uuid4())
        db.add(
            Case(
                id=case_id,
                case_no=c["case_no"],
                case_type="NORMAL",
                patent_category="INV",
                flow_dir="CN_DOMESTIC",
                client_id=client_ids[c["client_code"]],
                title_cn=c["title_cn"],
                title_en=c["title_en"],
                app_no=c["app_no"],
                status=c["status"],
                filing_date=c["filing_date"],
                recv_date=c["recv_date"],
            )
        )
        # Add one applicant per case
        db.add(
            T_CaseApplicant(
                id=str(uuid4()),
                case_id=case_id,
                seq=1,
                is_first=True,
                name_cn=c["applicant"],
            )
        )
        # Add one inventor per case
        db.add(
            T_CaseInventor(
                id=str(uuid4()),
                case_id=case_id,
                seq=1,
                name_cn=c["inventor"],
            )
        )

    db.commit()
    print(f"Created {len(v3_cases)} V3 test cases covering all 13 statuses")


def seed_masterdata_applicants(db: Session) -> None:
    """Seed core dev applicants with persisted applicant_type values. Idempotent."""
    seed_rows = [
        {
            "code": "DS-AP-001",
            "name_cn": "北京创新科技有限公司",
            "name_en": "Beijing Innovation Technology Co., Ltd.",
            "applicant_type": "ENTITY",
        },
        {
            "code": "DS-AP-002",
            "name_cn": "张三",
            "name_en": "Zhang San",
            "applicant_type": "INDIVIDUAL",
        },
    ]

    created = 0
    updated = 0
    for seed_row in seed_rows:
        existing = db.query(Applicant).filter(Applicant.code == seed_row["code"]).first()
        if not existing:
            db.add(
                Applicant(
                    id=str(uuid4()),
                    code=seed_row["code"],
                    name_cn=seed_row["name_cn"],
                    name_en=seed_row["name_en"],
                    applicant_type=seed_row["applicant_type"],
                    is_active=True,
                )
            )
            created += 1
            continue

        if existing.applicant_type != seed_row["applicant_type"]:
            existing.applicant_type = seed_row["applicant_type"]
            updated += 1

    db.commit()
    if created or updated:
        print(f"Seeded {created} dev applicants and updated {updated} applicant types")
    else:
        print("Dev applicants already seeded, skipping")


def seed_task_templates(db: Session) -> None:
    """Seed starter task templates. Idempotent."""
    templates = [
        {
            "code": "OA_REPLY",
            "name": "OA答复期限",
            "add_days": 120,
            "inner_offset_days": 14,
            "description": "审查意见通知书答复期限自动任务",
        },
        {
            "code": "OA_REPLY_SUBSEQUENT",
            "name": "后续审查意见答复期限",
            "add_days": None,
            "inner_offset_days": None,
            "description": "第二次及以后审查意见答复任务；截止日必须使用官文载明的明确期限",
        },
        {
            "code": "GRANT_FEE",
            "name": "授权登记费",
            "add_days": 60,
            "inner_offset_days": 7,
            "description": "授权登记费缴纳期限自动任务",
        },
    ]
    created = 0
    for t in templates:
        existing = db.query(TaskTemplate).filter(TaskTemplate.code == t["code"]).first()
        if not existing:
            db.add(
                TaskTemplate(
                    id=str(uuid4()),
                    code=t["code"],
                    name=t["name"],
                    add_days=t["add_days"],
                    inner_offset_days=t["inner_offset_days"],
                    description=t["description"],
                )
            )
            created += 1
    db.commit()
    if created:
        print(f"Created {created} task templates")
    else:
        print("Task templates already exist, skipping")


def seed_doc_templates(db: Session) -> None:
    """Seed default doc templates. Idempotent."""
    templates = [
        {
            "code": "OA_IN",
            "name": "审查意见通知书（收文）",
            "direction": "IN",
            "need_reply": True,
            "deadline_template_code": "OA_REPLY",
            "status_effect": "OA1",
            "status_restore": None,
            "fee_draft_type": None,
            "fee_item_list": None,
            "reply_to_template_code": None,
            "input_fields": None,
        },
        {
            "code": "OA_OUT",
            "name": "审查意见答复书（发文）",
            "direction": "OUT",
            "need_reply": False,
            "deadline_template_code": None,
            "status_effect": None,
            "status_restore": None,
            "fee_draft_type": None,
            "fee_item_list": None,
            "reply_to_template_code": "OA_IN",
            "input_fields": None,
        },
        {
            "code": "ACCEPTANCE_NOTICE",
            "name": "受理通知书",
            "direction": "IN",
            "need_reply": False,
            "deadline_template_code": None,
            "status_effect": "ACCEPTED",
            "status_restore": None,
            "fee_draft_type": None,
            "fee_item_list": None,
            "reply_to_template_code": None,
            "input_fields": None,
        },
        {
            "code": "GRANT_NOTICE",
            "name": "授权通知书",
            "direction": "IN",
            "need_reply": False,
            "deadline_template_code": None,
            "status_effect": "GRANT_PENDING",
            "status_restore": None,
            "fee_draft_type": "GRANT_FEE",
            "fee_item_list": None,
            "reply_to_template_code": None,
            "input_fields": None,
        },
        {
            "code": "CLIENT_IN",
            "name": "客户来函",
            "direction": "IN",
            "need_reply": False,
            "deadline_template_code": None,
            "status_effect": None,
            "status_restore": None,
            "fee_draft_type": None,
            "fee_item_list": None,
            "reply_to_template_code": None,
            "input_fields": None,
        },
    ]
    created = 0
    for t in templates:
        existing = db.query(DocTemplate).filter(DocTemplate.code == t["code"]).first()
        if not existing:
            db.add(DocTemplate(id=str(uuid4()), **t))
            created += 1
    official_notice_catalog_changed = seed_fee_reduction_approval_official_notice_catalog(db)
    official_letter_out_changed = seed_official_letter_out_catalog(db)
    grant_fee_notice_changed = seed_grant_fee_notice_template_source(db)
    format_letter_mapping_changed = seed_format_letter_mappings(db)
    db.commit()
    if (
        created
        or official_notice_catalog_changed
        or official_letter_out_changed
        or grant_fee_notice_changed
        or format_letter_mapping_changed
    ):
        print(
            "Created/updated "
            f"{created} doc templates, "
            f"{official_notice_catalog_changed} official notice catalog entries, "
            f"{official_letter_out_changed} official letter-out catalog entries, "
            f"{format_letter_mapping_changed} format letter mappings, "
            "and grant fee notice source"
        )
    else:
        print("Doc templates already exist, skipping")


def seed_grant_fee_notice_template_source(db: Session) -> bool:
    """Ensure grant-fee notice rendering has the local demo template source it requires."""
    _ensure_grant_fee_notice_docx_template()
    changed = False

    doc_template = (
        db.query(DocTemplate).filter(DocTemplate.code == GRANT_FEE_NOTICE_TEMPLATE_CODE).first()
    )
    doc_values = {
        "code": GRANT_FEE_NOTICE_TEMPLATE_CODE,
        "name": "授权费通知函",
        "direction": "OUT",
        "need_reply": False,
        "deadline_template_code": None,
        "status_effect": None,
        "status_restore": None,
        "fee_draft_type": None,
        "fee_item_list": None,
        "reply_to_template_code": None,
        "input_fields": None,
        "enabled": True,
    }
    if doc_template is None:
        db.add(DocTemplate(id=str(uuid4()), **doc_values))
        changed = True
    else:
        for field, value in doc_values.items():
            if getattr(doc_template, field) != value:
                setattr(doc_template, field, value)
                changed = True

    sources = (
        db.query(Template)
        .filter(
            Template.group == "DOC_TEMPLATE",
            Template.name == GRANT_FEE_NOTICE_TEMPLATE_CODE,
        )
        .order_by(Template.created_at.asc(), Template.id.asc())
        .all()
    )
    source_values = {
        "name": GRANT_FEE_NOTICE_TEMPLATE_CODE,
        "group": "DOC_TEMPLATE",
        "language": "zh-CN",
        "file_path": GRANT_FEE_NOTICE_TEMPLATE_PATH,
        "enabled": True,
    }
    source = sources[0] if sources else None
    for duplicate in sources[1:]:
        db.delete(duplicate)
        changed = True
    if source is None:
        db.add(Template(id=str(uuid4()), **source_values))
        changed = True
    else:
        for field, value in source_values.items():
            if getattr(source, field) != value:
                setattr(source, field, value)
                changed = True

    return changed


def _ensure_grant_fee_notice_docx_template() -> Path:
    template_path = BASE_DIR / "storage" / GRANT_FEE_NOTICE_TEMPLATE_PATH
    if template_path.exists():
        return template_path

    template_path.parent.mkdir(parents=True, exist_ok=True)
    document = DocxDocument()
    document.add_heading("授权费通知函", level=1)
    document.add_paragraph("案件编号：{{ case_no }}")
    document.add_paragraph("缴费期限：{{ grant_fee_task.due_date }}")
    document.add_paragraph(
        "官费金额：{{ grant_fee_task.gov_fee_amt }} {{ grant_fee_task.currency }}"
    )
    document.add_paragraph(
        "服务费金额：{{ grant_fee_task.service_fee_amt }} {{ grant_fee_task.currency }}"
    )
    document.add_paragraph("请根据授权通知要求确认缴费指示。")
    document.save(template_path)
    return template_path


class FormatLetterCatalogEntry(NamedTuple):
    code: str
    customer_format_letter_name: str
    official_doc_name_pattern: str
    source_path: str
    source_sha256: str
    template_version_id: str
    template_path: str
    expected_variables: frozenset[str]


FORMAT_LETTER_DATASET_ID = "FPMS-FORMAT-LETTER-CUSTOMER-20260610-V1"
FORMAT_LETTER_TEMPLATE_GROUP = "FORMAT_LETTER"
# 客户命名规则（信函生成操作 P0004）：{案号}-给{申请人名称}的邮件.docx
FORMAT_LETTER_OUTPUT_NAME_RULE = "{case_no}-给{applicant_name}的邮件.docx"
_FORMAT_LETTER_COMMON_VARIABLES = frozenset(
    {
        "salutation_text",
        "client_reference_no",
        "case_no",
        "invention_title",
        "application_no",
        "filing_date_text",
        "applicant_names_text",
        "source_notice_name",
    }
)


def _format_letter_variables(*extra: str) -> frozenset[str]:
    return _FORMAT_LETTER_COMMON_VARIABLES | frozenset(extra)


# 客户"国内客户天下先格式函对应官文"8 行（信函生成操作 P0007 TABLE 001）。
FORMAT_LETTER_MAPPING_CATALOG: tuple[FormatLetterCatalogEntry, ...] = (
    FormatLetterCatalogEntry(
        code="FORMAT_LETTER_001",
        customer_format_letter_name="官文转发-国内客户-驳回通知",
        official_doc_name_pattern="驳回决定",
        source_path="docs/postdemo/文件样例及模版/常用邮件模板/驳回通知.doc",
        source_sha256=("4f8f24d83bb3ca84f4663a0c46a1a84fa060ceb5881d2b9d3001fd074e81b4f2"),
        template_version_id="FPMS-FORMAT-LETTER-CUSTOMER-20260610-V1-001",
        template_path="templates/format_letters/format_letter_001.docx",
        expected_variables=_format_letter_variables("deadline_text"),
    ),
    FormatLetterCatalogEntry(
        code="FORMAT_LETTER_002",
        customer_format_letter_name="官文转发-国内客户-初审合格",
        official_doc_name_pattern="初步审查合格",
        source_path="docs/postdemo/文件样例及模版/常用邮件模板/初审合格.doc",
        source_sha256=("979713345eb2d5d8f4ee02421ebf5de6936b9f4418648dcffbe7a71f4cd62724"),
        template_version_id="FPMS-FORMAT-LETTER-CUSTOMER-20260610-V1-002",
        template_path="templates/format_letters/format_letter_002.docx",
        expected_variables=_format_letter_variables(),
    ),
    FormatLetterCatalogEntry(
        code="FORMAT_LETTER_003",
        customer_format_letter_name="官文转发-国内客户-公布通知",
        official_doc_name_pattern="公布通知书",
        source_path="docs/postdemo/文件样例及模版/常用邮件模板/公布通知.doc",
        source_sha256=("01502d4d3329adff358b3dfa0f995c5bcebb831b02155524140fbdf4995a81da"),
        template_version_id="FPMS-FORMAT-LETTER-CUSTOMER-20260610-V1-003",
        template_path="templates/format_letters/format_letter_003.docx",
        expected_variables=_format_letter_variables("publication_date_text", "deadline_text"),
    ),
    FormatLetterCatalogEntry(
        code="FORMAT_LETTER_004",
        customer_format_letter_name="官文转发-国内客户-实审通知",
        official_doc_name_pattern="进入实审通知",
        source_path="docs/postdemo/文件样例及模版/常用邮件模板/实审通知.doc",
        source_sha256=("dec38f3f2999b35c39b6d9cfa9f204bd0c930132da750170b1cdf90bd4666c00"),
        template_version_id="FPMS-FORMAT-LETTER-CUSTOMER-20260610-V1-004",
        template_path="templates/format_letters/format_letter_004.docx",
        expected_variables=_format_letter_variables(),
    ),
    FormatLetterCatalogEntry(
        code="FORMAT_LETTER_005",
        customer_format_letter_name="官文转发-国内客户-受通",
        official_doc_name_pattern="受理通知-电子",
        source_path="docs/postdemo/文件样例及模版/常用邮件模板/受通.doc",
        source_sha256=("acc06a13d4b09349d2eb81fadd31509ca7e260047df21ab38d65621d4607fff0"),
        template_version_id="FPMS-FORMAT-LETTER-CUSTOMER-20260610-V1-005",
        template_path="templates/format_letters/format_letter_005.docx",
        expected_variables=_format_letter_variables("inventor_names_text"),
    ),
    FormatLetterCatalogEntry(
        code="FORMAT_LETTER_006",
        customer_format_letter_name="官文转发-国内客户-授权通知",
        official_doc_name_pattern="授权通知书-电子",
        source_path="docs/postdemo/文件样例及模版/常用邮件模板/授权通知.doc",
        source_sha256=("83c14cf4b6514bee2c3f084cb13d2fb66f067dbeb8260de659f9ffb18e542974"),
        template_version_id="FPMS-FORMAT-LETTER-CUSTOMER-20260610-V1-006",
        template_path="templates/format_letters/format_letter_006.docx",
        expected_variables=_format_letter_variables("deadline_text", "amount_lines_text"),
    ),
    FormatLetterCatalogEntry(
        code="FORMAT_LETTER_007",
        customer_format_letter_name="官文转发-国内客户-一通",
        official_doc_name_pattern="第一次审查意见通知书",
        source_path="docs/postdemo/文件样例及模版/常用邮件模板/审查意见或复审通知.doc",
        source_sha256=("59322585ca96505fd1f38536ffa02e7afd3825df74c750e5256708b3869484cc"),
        template_version_id="FPMS-FORMAT-LETTER-CUSTOMER-20260610-V1-007",
        template_path="templates/format_letters/format_letter_007.docx",
        expected_variables=_format_letter_variables("deadline_text", "notice_variant_code"),
    ),
    FormatLetterCatalogEntry(
        code="FORMAT_LETTER_008",
        customer_format_letter_name="官文转发-专利证书",
        official_doc_name_pattern="专利证书",
        source_path="docs/postdemo/文件样例及模版/常用邮件模板/专利证书.doc",
        source_sha256=("5e9cf0638f02f98784c177e586f1127dc9fc0cc4e902596f61b368bb777c8be6"),
        template_version_id="FPMS-FORMAT-LETTER-CUSTOMER-20260610-V1-008",
        template_path="templates/format_letters/format_letter_008.docx",
        expected_variables=_format_letter_variables("inventor_names_text"),
    ),
)


def _validate_format_letter_templates() -> None:
    for entry in FORMAT_LETTER_MAPPING_CATALOG:
        template_path = BASE_DIR / "storage" / entry.template_path
        if not template_path.is_file():
            raise RuntimeError(f"FORMAT_LETTER_TEMPLATE_MISSING:{entry.code}")
        try:
            with zipfile.ZipFile(template_path) as package:
                package_xml = "\n".join(
                    package.read(name).decode("utf-8", errors="ignore")
                    for name in package.namelist()
                    if name.endswith(".xml")
                )
            variables = frozenset(
                DocxTemplate(str(template_path)).get_undeclared_template_variables()
            )
            document = DocxDocument(template_path)
        except Exception as exc:
            raise RuntimeError(f"FORMAT_LETTER_TEMPLATE_INVALID:{entry.code}") from exc

        visible_text = "\n".join(
            [paragraph.text for paragraph in document.paragraphs]
            + [
                paragraph.text
                for table in document.tables
                for row in table.rows
                for cell in row.cells
                for paragraph in cell.paragraphs
            ]
        )
        if (
            "MERGEFIELD" in package_xml
            or "请查收所附官方文件，并按期限完成后续事项" in visible_text
        ):
            raise RuntimeError(f"FORMAT_LETTER_TEMPLATE_PLACEHOLDER_REMAINS:{entry.code}")
        if variables != entry.expected_variables:
            raise RuntimeError(f"FORMAT_LETTER_TEMPLATE_VARIABLES_MISMATCH:{entry.code}")


def seed_format_letter_mappings(db: Session) -> int:
    """Seed the customer's 8-row official-notice -> format-letter mapping. Idempotent."""
    from app.modules.templates.models import FormatLetterMapping  # noqa: PLC0415

    _validate_format_letter_templates()
    resolved_carriers = []
    for entry in FORMAT_LETTER_MAPPING_CATALOG:
        mappings = (
            db.query(FormatLetterMapping)
            .filter(FormatLetterMapping.format_letter_template_code == entry.code)
            .order_by(FormatLetterMapping.created_at.asc(), FormatLetterMapping.id.asc())
            .all()
        )
        if len(mappings) > 1:
            raise RuntimeError(f"FORMAT_LETTER_MAPPING_AMBIGUOUS:{entry.code}")
        mapping = mappings[0] if mappings else None

        stable_templates = (
            db.query(Template)
            .filter(
                Template.group == FORMAT_LETTER_TEMPLATE_GROUP,
                Template.name == entry.code,
            )
            .order_by(Template.created_at.asc(), Template.id.asc())
            .all()
        )
        if len(stable_templates) > 1:
            raise RuntimeError(f"FORMAT_LETTER_TEMPLATE_ROW_AMBIGUOUS:{entry.code}")

        template = None
        if mapping is not None and mapping.format_letter_template_id is not None:
            linked_template = db.get(Template, mapping.format_letter_template_id)
            if (
                linked_template is None
                or linked_template.name != entry.code
                or linked_template.group != FORMAT_LETTER_TEMPLATE_GROUP
                or (stable_templates and stable_templates[0].id != linked_template.id)
            ):
                raise RuntimeError(f"FORMAT_LETTER_TEMPLATE_ROW_AMBIGUOUS:{entry.code}")
            template = linked_template
        elif stable_templates:
            template = stable_templates[0]
        resolved_carriers.append((entry, mapping, template))

    changed = 0
    for entry, mapping, template in resolved_carriers:
        template_values = {
            "name": entry.code,
            "group": FORMAT_LETTER_TEMPLATE_GROUP,
            "language": "zh-CN",
            "file_path": entry.template_path,
            "enabled": True,
        }
        if template is None:
            template = Template(id=str(uuid4()), **template_values)
            db.add(template)
            db.flush()
            changed += 1
        else:
            template_changed = False
            for field, value in template_values.items():
                if getattr(template, field) != value:
                    setattr(template, field, value)
                    template_changed = True
            if template_changed:
                changed += 1

        provenance = json.dumps(
            {
                "customer_format_letter_name": entry.customer_format_letter_name,
                "dataset_id": FORMAT_LETTER_DATASET_ID,
                "source_path": entry.source_path,
                "source_sha256": f"sha256:{entry.source_sha256}",
                "template_version_id": entry.template_version_id,
            },
            ensure_ascii=False,
            sort_keys=True,
            separators=(",", ":"),
        )
        mapping_values = {
            "official_doc_template_id": None,
            "official_doc_template_code": None,
            "official_doc_name_pattern": entry.official_doc_name_pattern,
            "format_letter_template_id": template.id,
            "format_letter_template_code": entry.code,
            "output_name_rule": FORMAT_LETTER_OUTPUT_NAME_RULE,
            "salutation_rule_code": "PRIMARY_CONTACT_TITLE",
            "contact_rule_code": "CLIENT_PRIMARY_CONTACT",
            "enabled": True,
            "remark": provenance,
        }
        if mapping is None:
            db.add(FormatLetterMapping(id=str(uuid4()), **mapping_values))
            changed += 1
            continue
        mapping_changed = False
        for field, value in mapping_values.items():
            if getattr(mapping, field) != value:
                setattr(mapping, field, value)
                mapping_changed = True
        if mapping_changed:
            changed += 1
    return changed


def seed_official_fee_rate_catalog(db: Session) -> None:
    """Seed post-demo official fee rate parameters. Idempotent."""
    created = 0
    updated = 0
    for values in _official_fee_rate_catalog():
        fee_code = values["fee_code"]
        existing = db.query(FeeRate).filter(FeeRate.fee_code == fee_code).one_or_none()
        if existing is None:
            db.add(FeeRate(id=str(uuid4()), **values))
            created += 1
            continue

        changed = False
        for field, value in values.items():
            if getattr(existing, field) != value:
                setattr(existing, field, value)
                changed = True
        if changed:
            updated += 1

    db.commit()
    if created or updated:
        print(f"Seeded {created} official fee rates and updated {updated}")
    else:
        print("Official fee rate catalog already seeded, skipping")


def seed_system_params(db: Session) -> None:
    """Seed default system parameters. Idempotent."""
    defaults = [
        {"param_key": "case_no_prefix", "param_value": "CN", "description": "案号前缀"},
        {"param_key": "default_currency", "param_value": "CNY", "description": "默认币种"},
        {"param_key": "bill_no_prefix", "param_value": "INV", "description": "账单编号前缀"},
        {
            "param_key": "task_sheet_template_path",
            "param_value": "templates/task_sheet.docx",
            "description": "任务单模板路径",
        },
    ]
    created = 0
    for d in defaults:
        existing = db.query(SystemParam).filter(SystemParam.param_key == d["param_key"]).first()
        if not existing:
            db.add(SystemParam(**d))
            created += 1
    db.commit()
    if created:
        print(f"Created {created} system parameters")
    else:
        print("System parameters already exist, skipping")


def main() -> None:
    """Run all dev seeds."""
    with SessionLocal() as db:
        print("Seeding default roles and permissions...")
        seed_default_roles_perms(db)
        print("✓ Roles and permissions seeded")

        print("Seeding admin user...")
        seed_admin_user(db)
        print("✓ Admin user seeded")

        print("Seeding V3 workflow test cases...")
        seed_v3_cases(db)
        print("✓ V3 test cases seeded")

        print("Seeding masterdata applicants...")
        seed_masterdata_applicants(db)
        print("✓ Masterdata applicants seeded")

        print("Seeding task templates...")
        seed_task_templates(db)
        print("✓ Task templates seeded")

        print("Seeding doc templates...")
        seed_doc_templates(db)
        print("✓ Doc templates seeded")

        print("Seeding system parameters...")
        seed_system_params(db)
        print("✓ System parameters seeded")

        print("Seeding official fee rate catalog...")
        seed_official_fee_rate_catalog(db)
        print("✓ Official fee rate catalog seeded")

    print("\n✅ Development database seeded successfully!")
    print("   Login: username='admin', password='admin123'")


if __name__ == "__main__":
    try:
        main()
    except SystemExit:
        raise
    except Exception as exc:  # noqa: BLE001
        print(f"Seed failed: {exc}")
        sys.exit(1)
